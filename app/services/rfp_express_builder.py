"""
RFP Kit Express Builder — SGD 129
==================================
Generates a 2-page evidence certificate for vendors responding to GeBIZ RFPs.

Flow:
  1. Derive vendor context from DB (company, UEN, sector, score)
  2. Generate 5 essential RFP Q&A answers via BooppaAIService
  3. Build PDF certificate via PDFService
  4. Upload PDF to S3 via S3Service
  5. Send delivery email via RFPExpressEmailer
  6. Return download URL + metadata

RFP Kit Complete (SGD 499) follows the same flow with 15 questions and
an editable DOCX — handled by a separate builder.
"""
from __future__ import annotations


import logging
import re
import uuid
from datetime import datetime, timedelta, timezone
from typing import Any, Dict, Optional

logger = logging.getLogger(__name__)


def _same_site(href_netloc: str, base_netloc: str) -> bool:
    """True when two hostnames belong to the same registrable site.

    Used to keep a scraped "privacy policy" link honest: a vendor page routinely
    links out to *other* privacy policies (Google reCAPTCHA's, a CDN's, an
    analytics vendor's). Only a link on the vendor's own domain (or a subdomain)
    may be presented as *their* published policy.
    """
    h = (href_netloc or "").lower().strip().removeprefix("www.")
    b = (base_netloc or "").lower().strip().removeprefix("www.")
    if not h or not b:
        return False
    return h == b or h.endswith("." + b) or b.endswith("." + h)


def _privacy_url_belongs_to_site(privacy_url: Optional[str], vendor_url: Optional[str]) -> bool:
    """Backstop before a privacy-policy URL is shown as the vendor's own.

    Applied at the point the URL enters the report context, so a foreign URL
    never reaches the PDF regardless of where it came from (a stale pre-guard
    cache entry, an AI-suggested link, any future code path) — not only when the
    live scraper first extracts it. Fixes the recurring "Privacy Policy:
    policies.google.com" leak on vendors that embed Google reCAPTCHA.
    """
    if not privacy_url or not vendor_url:
        return False
    from urllib.parse import urlparse
    return _same_site(urlparse(privacy_url).netloc, urlparse(vendor_url).netloc)


# ── Question sets ─────────────────────────────────────────────────────────────
# Express (5 questions) — core GeBIZ requirements
ESSENTIAL_QUESTIONS = [
    "data_policy",       # PDPA / data handling policy
    "dpo_appointed",     # DPO appointment status
    "security_measures", # Technical/organisational security controls
    "breach_history",    # Incident history (last 24 months)
    "third_party",       # Third-party vendor / sub-processor management
]

# Complete (15 questions) — full procurement evidence pack
COMPLETE_QUESTIONS = ESSENTIAL_QUESTIONS + [
    "iso_certifications",   # ISO 27001 / SOC 2 status
    "business_continuity",  # BCP / DR plan
    "staff_training",       # Security awareness training
    "access_controls",      # IAM and privileged access management
    "vulnerability_mgmt",   # Patch management and vulnerability scanning
    "encryption_standards", # Encryption algorithms and key management
    "audit_logging",        # Audit log retention and monitoring
    "incident_response",    # Incident response plan and contact
    "data_residency",       # Where data is stored (Singapore / overseas)
    "subcontracting",       # Subcontracting / offshoring policy
]

QUESTION_LABELS: dict[str, str] = {
    "data_policy":          "Do you have a PDPA data protection policy?",
    "dpo_appointed":        "Has a Data Protection Officer (DPO) been appointed?",
    "security_measures":    "What security measures are in place to protect personal data?",
    "breach_history":       "Have there been any data breaches in the past 24 months?",
    "third_party":          "How do you manage third-party vendors who handle personal data?",
    "iso_certifications":   "Does your organisation hold ISO 27001, SOC 2, or equivalent certification?",
    "business_continuity":  "Do you have a Business Continuity / Disaster Recovery plan?",
    "staff_training":       "How do you train staff on data protection and cybersecurity?",
    "access_controls":      "Describe your Identity and Access Management (IAM) controls.",
    "vulnerability_mgmt":   "How do you manage software vulnerabilities and patching?",
    "encryption_standards": "What encryption standards do you use for data at rest and in transit?",
    "audit_logging":        "How long are audit logs retained and how are they monitored?",
    "incident_response":    "Describe your incident response process and escalation path.",
    "data_residency":       "Where is data stored — Singapore, or overseas? What cross-border safeguards apply?",
    "subcontracting":       "Do you subcontract or offshore any processing involving personal data?",
}


class RFPExpressBuilder:
    """Generate RFP Kit Express package for a vendor."""

    def __init__(self, vendor_id: str, vendor_email: str, session_id: str | None = None):
        self.vendor_id    = vendor_id
        self.vendor_email = vendor_email
        self.session_id   = session_id
        # 4.12: Idempotent report_id — same session always produces the same ID (safe retries)
        if session_id:
            self.report_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"rfp:{session_id}"))
        else:
            self.report_id = str(uuid.uuid4())
        self.errors: list[str]    = []
        self.warnings: list[str]  = []
        self.used_template: bool  = False  # 4.3: track if AI failed and template was used
        self.generation_start     = datetime.now(timezone.utc)

    # ── Public entry point ────────────────────────────────────────────────────

    async def generate_express_package(
        self,
        vendor_url: str,
        company_name: str,
        rfp_details: Optional[Dict] = None,
        db=None,
        product_type: str = "rfp_express",
        allow_incomplete: bool = False,
    ) -> Dict[str, Any]:
        logger.info(f"RFP Kit Express: starting for {company_name} ({vendor_url})")

        questions = COMPLETE_QUESTIONS if product_type == "rfp_complete" else ESSENTIAL_QUESTIONS

        intake = (rfp_details or {}).get("intake", {})

        # 1. Gather vendor context (DB: UEN, sector, trust score, ACRA local match)
        vendor_ctx = self._build_vendor_context(company_name, vendor_url, db, intake=intake)

        # 1b. External evidence enrichment (parallel async calls)
        from app.services.dns_security import fetch_dns_security
        from app.services.gebiz_service import get_vendor_gebiz_history
        from app.services.onemap_service import fetch_onemap_location
        from app.services.evidence_enricher import (
            fetch_acra_status,
            fetch_pdpc_enforcement,
            fetch_ssl_grade,
            fetch_domain_reputation,
            fetch_hosting_signals,
            check_consistency,
        )
        from app.core.db import SessionLocal
        
        import asyncio as _asyncio
        uen = vendor_ctx.get("uen") or intake.get("uen")

        stated_hosting = intake.get("data_hosting") or intake.get("primary_cloud")
        (
            acra_live, pdpc_result, ssl_result, domain_rep, hosting_signals,
            dns_security
        ) = await _asyncio.gather(
            fetch_acra_status(uen, company_name),
            fetch_pdpc_enforcement(company_name, uen),
            fetch_ssl_grade(vendor_url),
            fetch_domain_reputation(vendor_url),
            fetch_hosting_signals(vendor_url, stated_hosting=stated_hosting),
            fetch_dns_security(vendor_url),
        )

        gebiz_history = {"checked": False, "total_awards": 0, "awards": []}
        try:
            if company_name:
                with SessionLocal() as db_session:
                    gebiz_history = await _asyncio.to_thread(get_vendor_gebiz_history, db_session, company_name)
        except Exception as e:
            logger.warning("GeBIZ enrichment failed: %s", e)
            
        onemap_location = {"checked": False, "found": False}
        try:
            postal = acra_live.get("postal_code")
            if postal:
                onemap_location = await fetch_onemap_location(str(postal))
        except Exception as e:
            logger.warning("OneMap enrichment failed: %s", e)

        # Merge live ACRA data into vendor context
        if acra_live.get("found"):
            self._merge_acra_into_ctx(vendor_ctx, acra_live)
            if acra_live.get("warning"):
                self.warnings.append(f"[ACRA] {acra_live['warning']}")
        if pdpc_result.get("warning"):
            self.warnings.append(f"[PDPC] {pdpc_result['warning']}")
        if ssl_result.get("warning"):
            self.warnings.append(f"[SSL] {ssl_result['warning']}")
        if domain_rep.get("warning"):
            self.warnings.append(f"[VT] {domain_rep['warning']}")
        if hosting_signals.get("mismatch_warning"):
            self.warnings.append(f"[Hosting] {hosting_signals['mismatch_warning']}")
        # Surface inferred provider in AI context
        if hosting_signals.get("inferred_provider"):
            vendor_ctx["inferred_hosting_provider"] = hosting_signals["inferred_provider"]
            vendor_ctx["inferred_hosting_region"] = hosting_signals.get("inferred_region")

        # 4.4: Consistency check + website scrape (single fetch, reused below)
        ws_data = await self._fetch_website_context(vendor_url)
        website_text = ws_data.get("text", "") if isinstance(ws_data, dict) else (ws_data or "")
        # Surface privacy policy URL in vendor_ctx for PDF/email — but only if it
        # is on the vendor's own domain (never Google reCAPTCHA's / a CDN's).
        if (
            isinstance(ws_data, dict)
            and ws_data.get("privacy_policy_url")
            and _privacy_url_belongs_to_site(ws_data["privacy_policy_url"], vendor_url)
        ):
            vendor_ctx.setdefault("privacy_policy_url", ws_data["privacy_policy_url"])
        # SPA warning
        if isinstance(ws_data, dict) and ws_data.get("spa_warning") and ws_data["spa_warning"] not in self.warnings:
            self.warnings.append(f"[SPA] {ws_data['spa_warning']}")

        # 2. Generate RFP Q&A answers via AI (enriched prompt) — pass ws_data so no re-fetch
        qa_answers = await self._generate_qa(
            vendor_ctx, rfp_details, questions,
            ssl_result=ssl_result, domain_rep=domain_rep,
            ws_data=ws_data,
            dns_security=dns_security,
            gebiz_history=gebiz_history, onemap_location=onemap_location
        )

        # 2a. Fill the placeholders the buyer ALREADY answered in their intake.
        # The anti-fabrication prompt makes the AI leave "[Verify: …]" markers
        # for facts it can't ground from the website — but when the buyer
        # supplied that exact fact in the intake (ISO cert number, BCP test
        # date, training cadence, sub-processors), the kit must use it instead
        # of telling the buyer to re-enter what they already gave us. (Forensic
        # audit: supplied facts were still shipping as "[Verify: …]".)
        qa_answers = self._apply_intake_substitutions(qa_answers, intake)

        # 2a-gate. HARD pre-delivery gate (audit fix). An RFP kit that still
        # contains ANY [Verify:] / [FILL IN] placeholder after intake substitution
        # is NOT delivered. GeBIZ-bound documents must be complete — a kit with
        # unfilled verification markers is unusable, and previously it shipped
        # anyway (forensic audit: 14-15 unfilled [Verify:] delivered, UEN "Not
        # provided", empty checklist → REJECT). Rather than build / anchor / upload
        # / email an unusable kit, we return a `blocked` result so the caller routes
        # the buyer back to the intake form to supply the missing facts, then
        # regenerates. A fuller intake both feeds the AI prompt (fewer placeholders
        # generated) and the substitution pass (placeholders filled).
        residual_placeholders = self._count_residual_placeholders(qa_answers)
        if residual_placeholders and allow_incomplete:
            # Admin test-checkout / test_simulation path: a thin or empty brief
            # would normally block delivery, but the end-to-end test must still
            # yield a kit. Degrade the hard gate to a warning and build/anchor/
            # deliver anyway, leaving the surviving placeholders in qa_answers.
            logger.warning(
                "[RFP] allow_incomplete=True — delivering %s with %d residual "
                "placeholder(s) (test path)", company_name, residual_placeholders,
            )
        elif residual_placeholders:
            missing_fields = self._residual_placeholder_details(qa_answers)
            logger.warning(
                "[RFP] BLOCKED delivery for %s — %d residual placeholder(s) remain "
                "(used_template=%s)", company_name, residual_placeholders, self.used_template,
            )
            return {
                "success": False,
                "blocked": True,
                "residual_placeholders": residual_placeholders,
                "used_template": self.used_template,
                "missing_fields": missing_fields,
                "qa_answers": qa_answers,
                "company_name": company_name,
                "product_type": product_type,
                "warnings": self.warnings,
            }

        # 2b. Post-AI validation pass — catch AI hallucinations against intake.
        # LLMs over-confidently fill in DPO names, ISO certifications, hosting
        # regions, etc. even when the intake declared otherwise. We surface
        # these on the result page so the buyer fixes them before submitting.
        # Merged into the consistency-check discrepancies below so they land in
        # the PDF warnings + frontend `discrepancies` array exactly once.
        ai_discrepancies = self._validate_answers_against_intake(qa_answers, intake)

        # Consistency check — intake vs external evidence
        discrepancies = check_consistency(intake, website_text, pdpc_result, domain_rep)
        discrepancies = (discrepancies or []) + ai_discrepancies
        if discrepancies:
            self.warnings.extend([f"[Discrepancy] {d}" for d in discrepancies])

        # Extract structured signals from the scraped website. Powers two
        # things: (a) AI prompt — verified facts the LLM can name without
        # tripping anti-fabrication rules; (b) per-answer verification source
        # so the result page can show "Verified on your website" instead of
        # the generic "AI-generated" label.
        from app.services.evidence_enricher import extract_website_signals
        website_signals = extract_website_signals(website_text)

        # Compute per-answer verification — {source, evidence} dict per Q key
        verification_map = self._compute_verification(
            intake=intake, vendor_ctx=vendor_ctx, website_signals=website_signals,
            ssl_result=ssl_result, domain_rep=domain_rep, acra_live=acra_live,
        )

        # 2.5. Anchor the content-bound evidence hash to the blockchain. Compute
        # it from the finished Q&A first so the SAME SHA-256 is anchored AND
        # printed in the PDF (independently verifiable on EvidenceAnchorV3).
        self.evidence_hash = self._compute_evidence_hash(company_name, qa_answers)
        tx_hash = await self._anchor_to_blockchain()

        # Compute coverage summary for Gap 3
        ran_sources = []
        if acra_live.get("found"): ran_sources.append("ACRA")
        if pdpc_result.get("checked") or pdpc_result.get("found"): ran_sources.append("PDPC")
        if ssl_result.get("checked") or ssl_result.get("grade"): ran_sources.append("SSL")
        if vendor_ctx.get("gebiz_supplier") or vendor_ctx.get("gebiz_contracts_count"): ran_sources.append("GeBIZ")
        
        has_uen = bool((intake and intake.get("uen")) or vendor_ctx.get("uen"))
        
        coverage_parts = []
        if ran_sources:
            coverage_parts.append(f"Verified against {', '.join(ran_sources)}.")
        if not has_uen:
            coverage_parts.append("ACRA/GeBIZ: not run — UEN not provided/found.")
            
        coverage_summary = " ".join(coverage_parts)

        # 3. Build PDF
        pdf_bytes = self._build_pdf(
            company_name, vendor_url, qa_answers, vendor_ctx, tx_hash, product_type,
            acra_live=acra_live, pdpc_result=pdpc_result, discrepancies=discrepancies,
            intake=intake, verification_map=verification_map, coverage_summary=coverage_summary,
            gebiz_history=gebiz_history, dns_security=dns_security, onemap_location=onemap_location,
        )

        # 4. Upload to S3
        download_url = await self._upload_pdf(pdf_bytes, product_type)

        # 4b. For Complete tier, also generate and upload DOCX
        docx_url = None
        declaration_url = None
        appendix_d_url = None
        if product_type == "rfp_complete":
            docx_bytes = self._build_docx(
                company_name, vendor_url, qa_answers, vendor_ctx, tx_hash, product_type,
                intake=intake, coverage_summary=coverage_summary
            )
            if docx_bytes:
                uploaded = await self._upload_docx(docx_bytes)
                if uploaded:
                    # Emit the STABLE re-presign endpoint, not the raw 7-day
                    # presigned URL — the latter dies after a week even though the
                    # S3 object persists, which is why the DOCX "went missing" from
                    # delivered kits. `uploaded` being truthy (upload succeeded) is
                    # what keeps the rfp_complete completeness gate honest.
                    from app.core.config import settings
                    _api_base = (settings.API_PUBLIC_BASE_URL or settings.VERIFY_BASE_URL).rstrip("/")
                    docx_url = f"{_api_base}/api/reports/{self.report_id}/rfp-docx"

            # 4b-ii. Supplier Compliance Declaration (Sprint 5c) — the third
            # output. A neutral, defensible alternative to the non-standard
            # "GeBIZ Appendix D": consolidates the supplier declarations that
            # recur across SG government tenders, each tagged Verified vs
            # Client-Declared. Best-effort — never blocks delivery of the kit.
            try:
                from app.services.rfp_declaration_generator import build_supplier_declaration_pdf

                decl_score = pdpc_result.get("compliance_score") if isinstance(pdpc_result, dict) else None
                if decl_score is None:
                    # Re-fetch the score locally — if the PDPA scan finished while the RFP
                    # kit was building (they run concurrently), we want the final score here,
                    # not the 'Pending' state from 3 minutes ago when this task began.
                    from app.services.pdpa_findings import latest_pdpa_score
                    fresh_score = latest_pdpa_score(db, self.vendor_id)
                    # If `self.vendor_id` is an email, it gracefully handles it, but maybe we should use user_id if we have it:
                    if fresh_score is None and vendor_ctx.get("uen"):
                        # We don't have the user object here directly, so if fresh_score returns None
                        # fallback to vendor_ctx which might have it already
                        pass
                    
                    if fresh_score is not None:
                        decl_score = fresh_score
                    else:
                        decl_score = vendor_ctx.get("compliance_score")
                declaration_bytes = build_supplier_declaration_pdf(
                    company_name=company_name,
                    vendor_ctx=vendor_ctx,
                    intake=intake,
                    verification_map=verification_map,
                    acra_live=acra_live,
                    pdpc_result=pdpc_result,
                    compliance_score=decl_score,
                    tx_hash=tx_hash,
                    report_id=self.report_id,
                )
                if declaration_bytes:
                    declaration_url = await self._upload_declaration(declaration_bytes)
            except Exception as decl_err:
                logger.warning(f"Supplier declaration generation failed (non-blocking): {decl_err}")
                self.warnings.append(f"Declaration error: {decl_err}")

            # 4b-iii. "Appendix D" data-protection appendix (best-effort generic).
            # Reproduces the kit's data-protection Q&A as a numbered D.1..D.n
            # template the bidder can renumber to match their specific ITT — a
            # usable answer to the (non-standard) "GeBIZ Appendix D" ask, with a
            # prominent template disclaimer. Best-effort — never blocks delivery.
            try:
                from app.services.rfp_appendix_d_generator import build_appendix_d_pdf

                apx_score = pdpc_result.get("compliance_score") if isinstance(pdpc_result, dict) else None
                if apx_score is None:
                    apx_score = vendor_ctx.get("compliance_score")
                qa_items = [
                    {
                        "question": self._q_label(k),
                        "answer": v,
                        "verified": (
                            (verification_map.get(k) or {}).get("source", "ai_drafted") != "ai_drafted"
                            and not self._PLACEHOLDER_RE.search(v or "")
                        ),
                        "evidence": (verification_map.get(k) or {}).get("evidence", []),
                    }
                    for k, v in qa_answers.items()
                ]
                appendix_d_bytes = build_appendix_d_pdf(
                    company_name=company_name,
                    qa_items=qa_items,
                    vendor_ctx=vendor_ctx,
                    intake=intake,
                    acra_live=acra_live,
                    compliance_score=apx_score,
                    tx_hash=tx_hash,
                    report_id=self.report_id,
                    coverage_summary=coverage_summary,
                )
                if appendix_d_bytes:
                    appendix_d_url = await self._upload_appendix_d(appendix_d_bytes)
            except Exception as apx_err:
                logger.warning(f"Appendix D generation failed (non-blocking): {apx_err}")
                self.warnings.append(f"Appendix D error: {apx_err}")

        # 4c. Write CertificateLog audit row (4.11)
        await self._write_certificate_log(pdf_bytes, download_url, db)

        # NOTE: the residual-placeholder check is now a HARD gate executed earlier
        # (section "2a-gate"): if any [Verify:]/[FILL IN] marker survives intake
        # substitution the kit is blocked before this point and never reaches
        # build/anchor/upload/email. By here, qa_answers is placeholder-free.

        # 5. Send email
        await self._send_email(company_name, download_url, product_type, docx_url=docx_url, declaration_url=declaration_url, appendix_d_url=appendix_d_url, pdf_bytes=pdf_bytes)

        elapsed = (datetime.now(timezone.utc) - self.generation_start).total_seconds()
        logger.info(f"RFP Kit Express complete in {elapsed:.1f}s for {company_name}")

        from app.core.config import settings
        explorer_base = settings.active_polygon_explorer_url.rstrip("/")

        is_complete = product_type == "rfp_complete"
        # Build labelled Q&A for frontend display, with verification source +
        # evidence per answer. `confidence` is derived from source for back-compat
        # ("fact" when intake/external/website verified, "generated" otherwise).
        qa_display = []
        for k, v in qa_answers.items():
            vinfo = verification_map.get(k) or {"source": "ai_drafted", "evidence": []}
            qa_display.append({
                "question": self._q_label(k),
                "answer": v,
                "confidence": "fact" if vinfo["source"] != "ai_drafted" else "generated",
                "verification": vinfo,
            })

        return {
            "success":        True,
            "product":        "rfp_kit_complete" if is_complete else "rfp_kit_express",
            "price":          "SGD 599" if is_complete else "SGD 249",
            "vendor_id":      self.vendor_id,
            "company_name":   company_name,
            "vendor_url":     vendor_url,
            "download_url":   download_url,
            "pdf_s3_key":     getattr(self, "pdf_s3_key", None),
            "docx_url":       docx_url,
            "declaration_url": declaration_url,
            "appendix_d_url": appendix_d_url,
            "qa_answers":     qa_display,
            "qa_answers_count": len(qa_display),
            "tx_hash":        tx_hash,
            "polygonscan_url": f"{explorer_base}/tx/{tx_hash}" if tx_hash else None,
            "network":        settings.active_polygon_network_name,
            "testnet_notice": settings.blockchain_notice,
            "upsell_available": not is_complete,
            "upsell_product": None if is_complete else "rfp_kit_complete",
            "upsell_price":   None if is_complete else "SGD 599",
            "errors":         self.errors,
            "warnings":       self.warnings,
            "answer_source":  "template" if self.used_template else "ai_grounded",
            "discrepancies":  discrepancies,
            "file_hash":      getattr(self, "evidence_hash", None),
            "data_sources": {
                "acra_verified":         acra_live.get("found", False),
                "acra_live":             acra_live.get("live"),
                "pdpc_checked":          pdpc_result.get("checked", False),
                "pdpc_flagged":          pdpc_result.get("found", False),
                "ssl_checked":           ssl_result.get("checked", False),
                "ssl_grade":             ssl_result.get("grade"),
                "vt_checked":            domain_rep.get("checked", False),
                "vt_flagged":            domain_rep.get("flagged", False),
                "vt_reputation":         domain_rep.get("reputation"),
                "hosting_checked":       hosting_signals.get("checked", False),
                "inferred_provider":     hosting_signals.get("inferred_provider"),
                "inferred_region":       hosting_signals.get("inferred_region"),
                "hosting_mismatch":      bool(hosting_signals.get("mismatch_warning")),
                "website_scraped":       bool(website_text),
                "privacy_policy_found":  bool(vendor_ctx.get("privacy_policy_url")),
                "gebiz_supplier":        vendor_ctx.get("gebiz_supplier", False),
                "gebiz_contracts":       vendor_ctx.get("gebiz_contracts_count", 0),
                "intake_supplied":       bool(intake),
                "ai_grounded":           not self.used_template,
            },
            "generated_at":   self.generation_start.isoformat(),
            "generation_time_seconds": elapsed,
            "expires_at":     (datetime.now(timezone.utc) + timedelta(days=7)).isoformat(),
        }

    # ── Step 1: vendor context ────────────────────────────────────────────────

    def _build_vendor_context(self, company_name: str, vendor_url: str, db, intake: dict | None = None) -> Dict:
        ctx: Dict[str, Any] = {
            "company_name": company_name,
            "vendor_url":   vendor_url,
            "uen":          intake.get("uen") if intake else None,
            "sector":       None,
            "trust_score":  None,
            "acra_name":    None,
            "acra_entity_type": None,
            "verification_depth": None,
            "gebiz_supplier": False,
            "gebiz_contracts_count": 0,
            "privacy_policy_url": None,
            "compliance_score": None,
        }
        if db is None:
            return ctx
        try:
            from app.core.models import User, VendorScore
            from app.core.models import VendorSector
            import re as _re
            # vendor_id may be an email (for anonymous/no-account purchases) or a UUID
            if _re.match(r'^[0-9a-f-]{36}$', self.vendor_id or '', _re.IGNORECASE):
                user = db.query(User).filter(User.id == self.vendor_id).first()
            else:
                user = db.query(User).filter(User.email == self.vendor_id).first()
            if user:
                ctx["uen"] = getattr(user, "uen", None)
                # Single source of truth with the Cover Sheet's PDPA score, so
                # the Supplier Declaration prints the same number (e.g. 66/100)
                # instead of "not available".
                try:
                    from app.services.pdpa_findings import latest_pdpa_score
                    ctx["compliance_score"] = latest_pdpa_score(db, user.id)
                except Exception as score_err:
                    logger.warning(f"PDPA score lookup failed for {self.vendor_id}: {score_err}")
            # VendorScore / VendorSector are keyed by UUID — skip if vendor_id is an email
            is_uuid = _re.match(r'^[0-9a-f-]{36}$', self.vendor_id or '', _re.IGNORECASE)
            if is_uuid:
                score = db.query(VendorScore).filter(VendorScore.vendor_id == self.vendor_id).first()
                if score:
                    ctx["trust_score"] = score.total_score
            sector_row = db.query(VendorSector).filter(
                VendorSector.vendor_id == (user.id if user else self.vendor_id)
            ).first() if is_uuid or user else None
            if sector_row:
                ctx["sector"] = sector_row.sector

            # ACRA lookup — enrich with registered company name and entity type
            uen_to_check = ctx.get("uen") or (intake.get("uen") if intake else None)
            if uen_to_check:
                try:
                    from app.core.models import MarketplaceVendor
                    acra_row = db.query(MarketplaceVendor).filter(
                        MarketplaceVendor.uen == uen_to_check
                    ).first()
                    if acra_row:
                        ctx["acra_name"] = getattr(acra_row, "name", None) or getattr(acra_row, "company_name", None)
                        ctx["acra_entity_type"] = getattr(acra_row, "entity_type", None)
                        if not ctx.get("sector"):
                            ctx["sector"] = getattr(acra_row, "sector", None)
                        logger.info(f"ACRA match for UEN {uen_to_check}: {ctx['acra_name']}")
                except Exception as acra_err:
                    logger.warning(f"ACRA lookup failed for UEN {uen_to_check}: {acra_err}")

            # Audit fix D: GeBIZ supplier history
            try:
                from app.core.models import DiscoveredVendor
                disc = db.query(DiscoveredVendor).filter(
                    DiscoveredVendor.uen == uen_to_check
                ).first()
                if disc:
                    ctx["gebiz_supplier"] = bool(getattr(disc, "gebiz_supplier", False))
                    ctx["gebiz_contracts_count"] = getattr(disc, "gebiz_contracts_count", 0) or 0
                    logger.info(
                        f"GeBIZ supplier found for UEN {uen_to_check}: "
                        f"supplier={ctx['gebiz_supplier']} contracts={ctx['gebiz_contracts_count']}"
                    )
            except Exception as gebiz_err:
                logger.warning(f"GeBIZ supplier lookup failed for UEN {uen_to_check}: {gebiz_err}")
        except Exception as e:
            logger.warning(f"Could not fetch vendor context for {self.vendor_id}: {e}")
        return ctx

    @staticmethod
    def _merge_acra_into_ctx(vendor_ctx: Dict, acra_live: Dict) -> None:
        """Fold a successful live ACRA lookup into the report context.

        Crucially back-fills the UEN: a name-only ACRA match resolves the entity
        and returns its UEN, and without carrying that through the kit printed
        "UEN: Not provided" and claimed "ACRA/GeBIZ: not run — UEN not found"
        even though ACRA had just confirmed the company. Existing context values
        win (an intake-supplied UEN is never overwritten)."""
        if not vendor_ctx.get("acra_name"):
            vendor_ctx["acra_name"] = acra_live.get("registered_name")
        if not vendor_ctx.get("acra_entity_type"):
            vendor_ctx["acra_entity_type"] = acra_live.get("entity_type")
        if not vendor_ctx.get("uen") and acra_live.get("uen"):
            vendor_ctx["uen"] = acra_live["uen"]
        vendor_ctx["acra_live"] = acra_live.get("live", True)
        vendor_ctx["acra_status"] = acra_live.get("entity_status")

    # ── Step 2: AI-generated Q&A ──────────────────────────────────────────────

    async def _fetch_website_context(self, vendor_url: str) -> dict:
        """Fetch and extract readable text from the vendor's website for AI grounding.
        Returns dict: {text, privacy_policy_url, is_spa, spa_warning}
        Results are cached in Redis for 24 hours to avoid re-scraping on retries."""
        import hashlib, httpx
        from urllib.parse import urlparse
        from app.core.cache import cache as cache_mod

        # v3: invalidates pre-same-site-guard scrapes that cached a foreign
        # (e.g. policies.google.com) privacy-policy URL.
        cache_key = cache_mod.cache_key(f"rfp_scrape_v3:{hashlib.md5(vendor_url.encode(), usedforsecurity=False).hexdigest()}")
        cached = cache_mod.get(cache_key)
        if cached and isinstance(cached, dict) and "text" in cached:
            return cached

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            BeautifulSoup = None

        import re as _re
        texts: list[str] = []
        privacy_policy_url: str | None = None
        is_spa = False

        base = vendor_url.rstrip("/")
        parsed_base = urlparse(vendor_url)

        pages = [base, base + "/about", base + "/privacy-policy", base + "/privacy"]
        headers = {"User-Agent": "Mozilla/5.0 (compatible; BooppaBot/1.0)"}

        async with httpx.AsyncClient(timeout=10, follow_redirects=True) as client:
            for page_url in pages:
                try:
                    resp = await client.get(page_url, headers=headers)
                    if resp.status_code != 200:
                        continue
                    raw_html = resp.text

                    # Audit fix E: SPA detection — thin visible text + script tags
                    stripped = _re.sub(r'<[^>]+>', ' ', raw_html)
                    stripped = _re.sub(r'\s+', ' ', stripped).strip()
                    if len(stripped) < 300 and '<script' in raw_html:
                        is_spa = True

                    if BeautifulSoup:
                        soup = BeautifulSoup(raw_html, "lxml")
                        # Remove boilerplate tags
                        for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
                            tag.decompose()
                        text = soup.get_text(separator=" ", strip=True)

                        # Audit fix C/F: extract privacy policy URL from links
                        if privacy_policy_url is None:
                            for a in soup.find_all("a", href=True):
                                href = a["href"]
                                if any(kw in href.lower() for kw in ["privacy", "pdpa", "data-protection"]):
                                    if href.startswith("http"):
                                        parsed_href = urlparse(href)
                                        if _same_site(parsed_href.netloc, parsed_base.netloc):
                                            privacy_policy_url = href
                                    elif href.startswith("/"):
                                        privacy_policy_url = (
                                            f"{parsed_base.scheme}://{parsed_base.netloc}{href}"
                                        )
                                    if privacy_policy_url:
                                        break
                    else:
                        text = stripped

                    texts.append(text[:1500])
                except Exception:
                    pass

        full_text = "\n\n---\n\n".join(texts)[:4000] if texts else ""
        spa_warning = (
            "Website appears to be a JavaScript SPA — scraped content may be incomplete. "
            "AI answers will be less specific."
            if is_spa else None
        )
        result = {
            "text": full_text,
            "privacy_policy_url": privacy_policy_url,
            "is_spa": is_spa,
            "spa_warning": spa_warning,
        }
        if full_text:
            cache_mod.set(cache_key, result, ttl=86400)  # 24 h
        return result

    async def _generate_qa(
        self,
        ctx: Dict,
        rfp_details: Optional[Dict],
        questions: list,
        ssl_result: Dict | None = None,
        domain_rep: Dict | None = None,
        ws_data: dict | None = None,
        dns_security: Dict | None = None,
        gebiz_history: Dict | None = None,
        onemap_location: Dict | None = None,
    ) -> Dict[str, str]:
        try:
            from app.services.booppa_ai_service import BooppaAIService
            ai = BooppaAIService()

            sector_hint = f" in the {ctx['sector']} sector" if ctx.get("sector") else ""
            rfp_hint    = f" The RFP is for: {rfp_details.get('description', '')}." if rfp_details else ""
            keys_list   = ", ".join(questions)

            # Buyer-supplied facts ground the answers in reality
            intake = (rfp_details or {}).get("intake", {})
            intake_lines = []
            fact_map = {
                "uen":                "Company UEN (Singapore registration number)",
                "description":        "What the company does",
                "dpo_appointed":      "DPO appointed (yes/no/in-progress)",
                "dpo_name":           "DPO full name",
                "dpo_email":          "DPO contact email",
                "dpo_registration":   "PDPC DPO registration number",
                "iso_status":         "ISO 27001 / SOC 2 certification status",
                "iso_cert_number":    "ISO 27001 certificate number",
                "iso_cert_expiry":    "ISO 27001 certificate expiry date",
                "data_hosting":       "Where data is hosted",
                "primary_cloud":      "Primary cloud provider",
                "cloud_region":       "Cloud region where data is stored",
                "breach_history":     "Data breaches in last 24 months",
                "bcp_last_tested":    "Date BCP/DR plan was last tested",
                "training_frequency": "How often staff security training is conducted",
                "key_processors":     "Key third-party data processors",
                "extra_notes":        "Additional notes from the vendor",
            }
            for key, label in fact_map.items():
                val = intake.get(key)
                if val and val not in ("unknown", ""):
                    intake_lines.append(f"- {label}: {val}")
            if ctx.get("uen") and not intake.get("uen"):
                intake_lines.append(f"- Company UEN: {ctx['uen']}")
            if ctx.get("acra_name"):
                intake_lines.append(f"- ACRA registered name: {ctx['acra_name']}")
            if ctx.get("acra_entity_type"):
                intake_lines.append(f"- Entity type: {ctx['acra_entity_type']}")
            if ctx.get("gebiz_supplier"):
                count = ctx.get("gebiz_contracts_count") or ""
                count_str = f" with {count} prior government contracts" if count else ""
                intake_lines.append(f"- GeBIZ registered supplier{count_str}")
            if ctx.get("privacy_policy_url"):
                intake_lines.append(f"- Published privacy policy URL: {ctx['privacy_policy_url']}")
            if ctx.get("inferred_hosting_provider"):
                region_str = f" (region: {ctx['inferred_hosting_region']})" if ctx.get("inferred_hosting_region") else ""
                intake_lines.append(
                    f"- Inferred hosting provider: {ctx['inferred_hosting_provider']}{region_str}"
                )

            # Audit fix: warn when ISO claimed but cert number missing (not independently verifiable)
            if intake.get("iso_status") == "certified" and not intake.get("iso_cert_number"):
                self.warnings.append(
                    "ISO 27001 declared certified but no certificate number provided — "
                    "evaluators cannot independently verify. Add iso_cert_number to intake."
                )

            facts_section = (
                "Known facts about this company (use these to make answers specific and accurate):\n"
                + "\n".join(intake_lines)
                if intake_lines else ""
            )

            # Use pre-fetched website data if provided, otherwise fetch now (fallback)
            if ws_data is None:
                ws_data = await self._fetch_website_context(ctx["vendor_url"])
            _ws_data = ws_data
            website_text = _ws_data.get("text", "") if isinstance(_ws_data, dict) else (_ws_data or "")
            if isinstance(_ws_data, dict) and _ws_data.get("spa_warning"):
                if _ws_data["spa_warning"] not in self.warnings:
                    self.warnings.append(f"[SPA] {_ws_data['spa_warning']}")
            # Inject privacy policy URL into facts if not already provided —
            # same-site backstop so a foreign URL never reaches the prompt/PDF.
            if (
                isinstance(_ws_data, dict)
                and _ws_data.get("privacy_policy_url")
                and not ctx.get("privacy_policy_url")
                and _privacy_url_belongs_to_site(_ws_data["privacy_policy_url"], ctx.get("vendor_url"))
            ):
                ctx["privacy_policy_url"] = _ws_data["privacy_policy_url"]
                intake_lines.append(f"- Published privacy policy URL: {ctx['privacy_policy_url']}")
                facts_section = (
                    "Known facts about this company (use these to make answers specific and accurate):\n"
                    + "\n".join(intake_lines)
                    if intake_lines else ""
                )
            website_section = (
                f"Website content extract (use to infer actual products/services/practices):\n{website_text}"
                if website_text else ""
            )

            # Inject SSL and VirusTotal real data as additional facts
            external_facts = []
            if ssl_result and ssl_result.get("checked") and ssl_result.get("grade"):
                protos = ", ".join(ssl_result.get("protocols", [])) or "unknown"
                external_facts.append(f"- SSL Labs grade: {ssl_result['grade']} (protocols: {protos})")
            if domain_rep and domain_rep.get("checked"):
                if domain_rep.get("flagged"):
                    malicious = domain_rep.get("malicious_votes", 0)
                    external_facts.append(
                        f"- VirusTotal domain check: flagged as malicious by {malicious} vendor(s) — "
                        "disclose proactively in security_measures answer"
                    )
                else:
                    external_facts.append(
                        "- External domain reputation: no public security incidents or malicious flags detected. "
                        "For breach_history answer, state: 'No data breaches have occurred to the knowledge of management. "
                        "No external security incidents have been publicly reported against the company\\'s domains.' "
                        "Do NOT mention VirusTotal or any scanning tool names in the answer."
                    )

            if external_facts:
                external_section = "Verified external data:\n" + "\n".join(external_facts)
                facts_section = "\n\n".join(filter(None, [facts_section, external_section]))
                
            # Additional new external enrichment signals
            enrich_facts = []
            if dns_security and dns_security.get("checked"):
                if dns_security.get("dmarc_record"):
                    enrich_facts.append(f"- DNS Security: DMARC is enforced with policy p={dns_security.get('dmarc_policy')}")
                elif dns_security.get("spf_record"):
                    enrich_facts.append("- DNS Security: SPF is configured, but DMARC is missing")
            
            if gebiz_history and gebiz_history.get("total_awards", 0) > 0:
                enrich_facts.append(f"- GeBIZ Government Tenders: Vendor has successfully been awarded {gebiz_history['total_awards']} past government tenders.")
                
            if onemap_location and onemap_location.get("found"):
                enrich_facts.append(f"- Physical Location: Registered office is located in {onemap_location.get('planning_area')} (verified via OneMap).")
                
            if enrich_facts:
                enrich_section = "Additional Regulatory & Security Signals:\n" + "\n".join(enrich_facts)
                facts_section = "\n\n".join(filter(None, [facts_section, enrich_section]))

            # Verified from the buyer's own website / privacy policy. These are
            # safe for the LLM to name in answers because they're already
            # public — saying "ISO 27001 certified" matches what the buyer's
            # customers can already see. Without this section, the
            # anti-fabrication prompt blocks the LLM from using these even
            # when they're true and publicly stated.
            try:
                from app.services.evidence_enricher import extract_website_signals as _ews
                ws_signals = _ews(website_text)
            except Exception:
                ws_signals = {"available": False}

            verified_site_facts: list[str] = []
            if ws_signals.get("iso_27001_mentioned"):
                yr = ws_signals.get("iso_27001_year")
                verified_site_facts.append(f"- ISO 27001{':' + yr if yr else ''} certification is publicly referenced on the company's own website")
            if ws_signals.get("iso_27017_mentioned"):
                verified_site_facts.append("- ISO 27017 referenced on the company's website")
            if ws_signals.get("iso_27018_mentioned"):
                verified_site_facts.append("- ISO 27018 referenced on the company's website")
            if ws_signals.get("iso_27701_mentioned"):
                verified_site_facts.append("- ISO 27701 referenced on the company's website")
            if ws_signals.get("soc_2_mentioned"):
                verified_site_facts.append("- SOC 2 referenced on the company's website")
            if ws_signals.get("pci_dss_mentioned"):
                verified_site_facts.append("- PCI DSS referenced on the company's website")
            if ws_signals.get("aes_mentioned"):
                verified_site_facts.append("- Encryption with AES is referenced on the company's website")
            if ws_signals.get("tls_mentioned"):
                verified_site_facts.append("- TLS is referenced on the company's website")
            if ws_signals.get("singapore_residency_mentioned"):
                verified_site_facts.append("- Singapore data residency is referenced on the company's website")
            if ws_signals.get("non_sg_regions_mentioned"):
                verified_site_facts.append(
                    f"- Non-Singapore regions referenced on the company's website: {', '.join(ws_signals['non_sg_regions_mentioned'])}"
                )
            for provider_key, label in [("aws_mentioned", "AWS"), ("azure_mentioned", "Azure"),
                                         ("gcp_mentioned", "GCP"), ("oci_mentioned", "Oracle Cloud")]:
                if ws_signals.get(provider_key):
                    verified_site_facts.append(f"- {label} referenced as a cloud provider on the company's website")
            if ws_signals.get("dpa_mentioned"):
                verified_site_facts.append("- Data Processing Agreement language is referenced on the company's website")
            if ws_signals.get("subprocessors_mentioned"):
                verified_site_facts.append("- Sub-processor policy referenced on the company's website")
            if ws_signals.get("breach_policy_mentioned"):
                verified_site_facts.append("- Breach/incident response policy referenced on the company's website")
            if ws_signals.get("retention_policy_mentioned"):
                verified_site_facts.append("- Data retention policy referenced on the company's website")

            if verified_site_facts:
                verified_section = (
                    "Verified from the company's own published website "
                    "(safe to name in answers — these are already public):\n"
                    + "\n".join(verified_site_facts)
                )
                facts_section = "\n\n".join(filter(None, [facts_section, verified_section]))

            context_block = "\n\n".join(filter(None, [facts_section, website_section]))

            # Anti-fabrication prompt. Buyers were getting answers full of
            # specifics the LLM invented — "ISO 27001:2022 certified", "AES-256",
            # "99.99% uptime", "patched within 24 hours", "data in Singapore on
            # AWS and Azure". If submitted unread to a GeBIZ tender, those are
            # false statements in a government procurement context. Three rules:
            #   1. Anchor every specific to a known fact (intake / website /
            #      external data). No fact → no specific.
            #   2. When a fact is missing, write a "[Verify: ...]" placeholder.
            #      The buyer fills it in. The reviewer SEES the placeholder.
            #   3. Answer the EXACT question asked (matters for windows like
            #      "24 months" — the AI was answering "36 months" on intake-empty
            #      runs because nothing constrained it).
            prompt = (
                f"You are drafting RFP compliance answers for {ctx['company_name']}"
                f"{sector_hint} (website: {ctx['vendor_url']}).{rfp_hint}\n\n"
                + (f"{context_block}\n\n" if context_block else "")
                + "Write concise, professional answers for a Singapore government procurement RFP. "
                "Each answer must be 1-3 sentences.\n\n"
                "ABSOLUTE RULES — do not break these even if the answer reads less polished:\n\n"
                "1. NEVER invent specifics. Do not name any of the following unless the exact value "
                "appears in the facts above:\n"
                "   • Certifications or standards (e.g. ISO 27001, ISO 27701, SOC 2 — including versions/years)\n"
                "   • Encryption algorithms or protocols (e.g. AES-256, TLS 1.2/1.3, RSA-2048)\n"
                "   • Specific timeframes (e.g. '24 hours', 'within 7 days', 'weekly', 'quarterly')\n"
                "   • Uptime SLAs or percentages (e.g. '99.99%')\n"
                "   • Cloud providers, regions, or vendor names (e.g. AWS, Azure, GCP, Singapore region)\n"
                "   • Specific personnel names, contact emails, or DPO names\n"
                "   • Retention periods, audit log durations, or breach-notification windows\n"
                "2. When a specific would normally go but the fact is missing, write a bracketed "
                "placeholder the buyer must fill in: '[Verify: encryption standard]', "
                "'[Verify: ISO 27001 cert number and expiry]', '[Verify: BCP last test date]', "
                "'[Verify: SLA target]'. The placeholder is REQUIRED, not optional.\n"
                "3. Answer the EXACT question asked. If the question asks about '24 months', do NOT "
                "expand the window to 36 months. If the question asks 'Yes/No', begin with Yes or No.\n"
                "4. Do NOT cite legal obligations with specific numbers unless they are correct under "
                "Singapore PDPA. Breach notification to PDPC is 'within 3 calendar days' (PDPA §26D) — "
                "do NOT write '1 hour' or 'immediately'. The recipient is PDPC, not 'the government'.\n"
                "5. Do NOT speculate about technical architecture details (blockchain nodes, smart "
                "contract audits, threat monitoring) unless they appear verbatim in the facts above.\n\n"
                f"Return ONLY a JSON object with these exact keys:\n{keys_list}."
            )

            response = await ai._call_deepseek([{"role": "user", "content": prompt}])

            import json, re
            # Extract JSON block from AI response
            if response and isinstance(response, str):
                match = re.search(r'\{.*\}', response, re.DOTALL)
                if match:
                    ai_answers = json.loads(match.group())
                    if isinstance(ai_answers, dict):
                        # Backfill any question the model omitted with its template
                        # answer, so a partial JSON response never silently drops a
                        # question (leaving a blank in the deliverable). Missing keys
                        # get the [FILL IN] template so the placeholder gate still
                        # forces the buyer to complete them.
                        missing = [q for q in questions
                                   if not str(ai_answers.get(q) or "").strip()]
                        if missing:
                            self.warnings.append(
                                f"AI omitted {len(missing)} answer(s); backfilled from template"
                            )
                            tmpl = self._template_qa(ctx, missing)
                            for q in missing:
                                if q in tmpl:
                                    ai_answers[q] = tmpl[q]
                        return ai_answers
        except Exception as e:
            logger.warning(f"AI Q&A generation failed, using template fallback: {e}")
            self.warnings.append("AI Q&A used template fallback")

        # 4.3: Fallback — mark so PDF and result carry a disclosure
        self.used_template = True
        return self._template_qa(ctx, questions)

    # Audit fix A: per-answer prefix so evaluators know answers are templated.
    # Audit fix B (2026): template answers no longer assert specific facts the
    # buyer may not actually meet (e.g. "AES-256", "Singapore residency",
    # "RTO of 4 hours") — these turned a "starter draft" into a perjury risk
    # if a rushed buyer submitted unread to a GeBIZ tender. Specifics are now
    # `___ [FILL IN] ___` placeholders the buyer MUST replace before submitting.
    _TEMPLATE_PREFIX = "[Standard template — review and replace [FILL IN] fields before submission] "
    _FI = "___ [FILL IN] ___"

    def _template_qa(self, ctx: Dict, questions: list) -> Dict[str, str]:
        name = ctx["company_name"]
        url  = ctx["vendor_url"]
        p    = self._TEMPLATE_PREFIX
        fi   = self._FI
        all_answers = {
            "data_policy":          f"{p}{name} maintains a PDPA-compliant Personal Data Protection Policy, accessible at {url}. Personal data is collected with consent and retained only for its stated purpose.",
            "dpo_appointed":        f"{p}DPO appointment status: {fi} (Yes/No/In progress). If appointed, name: {fi}, email: {fi}, PDPC registration: {fi}.",
            "security_measures":    f"{p}{name} implements technical and organisational security controls including {fi} (e.g. access controls, encryption in transit, staff training). See encryption_standards and access_controls for specifics.",
            "breach_history":       f"{p}Breach history (last 24 months): {fi} (None / Yes — describe). Any PDPC-notifiable incidents must be disclosed here.",
            "third_party":          f"{p}{name} requires Data Processing Agreements (DPAs) before sharing personal data with sub-processors. Key processors in scope: {fi}.",
            "iso_certifications":   f"{p}ISO 27001 status: {fi} (Certified / Pursuing / None). If certified, certificate number: {fi}, expiry: {fi}. SOC 2 status: {fi}.",
            "business_continuity":  f"{p}{name} maintains a BCP/DR plan. Last tested: {fi}. RTO target: {fi}. RPO target: {fi}.",
            "staff_training":       f"{p}Security awareness training frequency: {fi} (e.g. annual, quarterly). New-hire training window: {fi}.",
            "access_controls":      f"{p}{name} enforces role-based access control with least-privilege principles. Privileged access review cadence: {fi}. MFA on privileged accounts: {fi} (Yes/No).",
            "vulnerability_mgmt":   f"{p}Patch SLA for critical vulnerabilities: {fi}. Vulnerability scan cadence: {fi}.",
            "encryption_standards": f"{p}Encryption at rest: {fi} (e.g. AES-256). Encryption in transit: {fi} (e.g. TLS 1.2+). Key management process: {fi}.",
            "audit_logging":        f"{p}Audit log retention period: {fi} months. Log monitoring/anomaly detection: {fi} (describe).",
            "incident_response":    f"{p}{name} maintains an Incident Response Plan. Internal notification window: {fi}. PDPC notification is made within 3 business days if the incident is PDPC-notifiable.",
            "data_residency":       f"{p}Primary data hosting region: {fi}. Cloud provider: {fi}. Cross-border transfers (if any) are governed by contractual clauses consistent with PDPA's Third Schedule.",
            "subcontracting":       f"{p}Subcontracting / offshoring of personal data processing: {fi} (None / Yes — list arrangements). Any engagements require prior written approval and binding DPAs.",
        }
        return {k: all_answers[k] for k in questions if k in all_answers}

    def _validate_answers_against_intake(self, qa_answers: Dict[str, str], intake: dict) -> list[str]:
        """Cross-check AI-generated answers against buyer-supplied intake facts.

        Two failure modes covered:
          (a) Contradiction — intake says X, answer says NOT X (e.g. intake
              dpo_appointed='no' but answer asserts a DPO exists).
          (b) Fabrication — intake has no value for X, but answer asserts a
              specific value anyway (e.g. intake silent on ISO 27001, answer
              claims "ISO 27001:2022 certified"). Submitting fabricated
              specifics to a GeBIZ tender is a legal risk under the Government
              Procurement Act.

        The original prompt is the first defence (see _generate_qa). This is the
        belt-and-suspenders pass that catches what the LLM lets through.

        Returns a list of human-readable discrepancy strings (may be empty).
        """
        if not isinstance(qa_answers, dict) or not isinstance(intake, dict):
            return []

        out: list[str] = []
        import re

        def lower(key: str) -> str:
            v = qa_answers.get(key)
            return v.lower() if isinstance(v, str) else ""

        # ── DPO ─────────────────────────────────────────────────────────────
        dpo_state = (intake.get("dpo_appointed") or "").lower()
        dpo_ans = lower("dpo_appointed")
        if dpo_state == "no" and dpo_ans:
            # Buyer declared no DPO, but answer claims one exists.
            if re.search(r"\b(has appointed|appointed a|dpo is|our dpo|the dpo)\b", dpo_ans) \
               and not re.search(r"\b(no dpo|has not appointed|not yet appointed|in progress|to be appointed)\b", dpo_ans):
                out.append(
                    "Intake declares no DPO appointed, but the dpo_appointed answer reads as if one exists. "
                    "Edit the answer to match the intake before submission."
                )

        # ── ISO 27001 ───────────────────────────────────────────────────────
        iso_state = (intake.get("iso_status") or "").lower()
        iso_ans = lower("iso_certifications")
        if iso_state in {"none", "pursuing", "in_progress", "unknown", ""} and iso_ans:
            if re.search(r"\bis (iso[\s-]?27001\s)?certified\b|\bholds (an? )?iso[\s-]?27001 certification\b|\biso[\s-]?27001 certified\b", iso_ans):
                out.append(
                    f"Intake declares ISO 27001 status as '{iso_state or 'unknown'}', but the answer asserts "
                    f"the company is certified. Correct the answer to match the intake."
                )
        if iso_state == "certified" and not intake.get("iso_cert_number"):
            out.append(
                "ISO 27001 declared certified in intake but no certificate number supplied — evaluators cannot verify. "
                "Add iso_cert_number to the intake or remove the certified claim."
            )

        # ── Breach history ──────────────────────────────────────────────────
        breach_state = (intake.get("breach_history") or "").lower()
        breach_ans = lower("breach_history")
        if breach_state in {"one", "multiple"} and breach_ans:
            if re.search(r"\bno (data )?breach(es)?\b|\bno security incidents?\b|\bnever experienced\b", breach_ans):
                out.append(
                    f"Intake declares {breach_state} breach(es) in the last 24 months, but the answer says none. "
                    "PDPC-notifiable breaches MUST be disclosed in tender responses."
                )

        # ── Data hosting region ────────────────────────────────────────────
        hosting = (intake.get("data_hosting") or "").lower()
        residency_ans = lower("data_residency")
        if hosting == "global" and residency_ans:
            if re.search(r"\b(all|only|exclusively).{0,30}\b(singapore|sg)\b", residency_ans):
                out.append(
                    "Intake declares global data hosting, but the data_residency answer claims Singapore-only. "
                    "Edit to disclose actual cross-border hosting."
                )
        if hosting in {"apac", "global"} and residency_ans:
            if re.search(r"\bdoes not (offshore|cross[- ]border transfer)\b", residency_ans):
                out.append(
                    f"Intake declares data_hosting='{hosting}' (non-SG), but the answer says no cross-border transfer. "
                    "Reconcile before submission."
                )

        # ── Fabricated specifics (fires when intake has no backing fact) ────
        # These patterns catch the LLM inventing certifications, encryption
        # standards, uptime SLAs, timeframes, cloud providers etc. when the
        # intake didn't supply them. We *don't* flag if the buyer actually
        # supplied a corresponding fact — they earned the right to that claim.

        def claims(key: str, pattern: str) -> bool:
            v = qa_answers.get(key)
            return bool(isinstance(v, str) and re.search(pattern, v, re.IGNORECASE))

        # ISO / SOC 2 / equivalent — flag specific certs/years when intake silent
        if iso_state in {"", "unknown", "none", "pursuing", "in_progress"}:
            for key in ("iso_certifications", "security_measures", "data_policy"):
                if claims(key, r"\biso[\s-]?2700[17](?::\s?\d{4})?\b|\bsoc[\s-]?2\b|\biso[\s-]?9001\b"):
                    out.append(
                        f"The '{key}' answer names a specific certification (ISO 27001 / SOC 2 / similar) but "
                        "the intake did not confirm any. Remove the certification claim or supply the "
                        "intake.iso_status + iso_cert_number to back it up."
                    )
                    break

        # Encryption algorithms — flag AES-N / TLS X.Y / RSA-N when intake silent
        if not (intake.get("encryption_standards_known") or intake.get("encryption_at_rest") or intake.get("encryption_in_transit")):
            for key in ("encryption_standards", "security_measures", "access_controls"):
                if claims(key, r"\baes[-\s]?(?:128|192|256)\b|\btls\s?1\.\d\b|\brsa[-\s]?\d{3,4}\b"):
                    out.append(
                        f"The '{key}' answer names a specific encryption standard (AES-N / TLS X.Y / RSA-N) "
                        "but the intake did not declare one. Confirm the actual standard used or replace "
                        "the specific with a [Verify: encryption standard] placeholder."
                    )
                    break

        # Uptime SLAs — flag 99.9...% when intake silent
        if not intake.get("uptime_sla"):
            for key in ("business_continuity", "security_measures", "incident_response"):
                if claims(key, r"\b99\.9{1,3}\s?%|\bfive[- ]nines?\b|\bfour[- ]nines?\b"):
                    out.append(
                        f"The '{key}' answer commits to a specific uptime SLA (99.9...% / five-nines) "
                        "but the intake did not declare one. Remove the percentage or supply the "
                        "actual SLA via intake."
                    )
                    break

        # Specific timeframes/cadences — flag "weekly / daily / monthly / N hours / N days"
        # in answers that should be backed by intake but weren't supplied.
        cadence_pat = r"\b(?:weekly|daily|monthly|quarterly|semi[- ]annually|annually|bi[- ]monthly)\b|\bwithin\s+\d+\s+(?:minute|hour|day|week|month)s?\b|\bevery\s+\d+\s+(?:minute|hour|day|week|month)s?\b"
        if not intake.get("training_frequency") and claims("staff_training", cadence_pat):
            out.append(
                "The 'staff_training' answer commits to a specific cadence but the intake did not "
                "declare one. Edit to your actual training schedule or replace with a [Verify: cadence] placeholder."
            )
        if not intake.get("bcp_last_tested") and claims("business_continuity", r"\btested\s+(?:weekly|daily|monthly|quarterly|semi[- ]annually|annually)\b"):
            out.append(
                "The 'business_continuity' answer commits to a BCP test cadence but the intake did "
                "not declare when the BCP was last tested. Confirm the actual cadence or use a placeholder."
            )
        # Vulnerability management — patch-SLA and scan-cadence specifics
        if claims("vulnerability_mgmt", r"\bpatched?\s+within\s+\d+\s+(?:hour|day)s?\b|\b(?:weekly|daily|monthly|quarterly)\s+(?:vulnerability scans?|penetration tests?)\b|\bwithin\s+\d+\s+(?:hour|day)s?\s+of\b"):
            out.append(
                "The 'vulnerability_mgmt' answer names a specific patch SLA or scan cadence. "
                "Confirm against your actual policy or replace with [Verify: patch SLA] / "
                "[Verify: scan cadence] placeholders."
            )
        # Audit log retention — flag specific N-year / N-month retention
        if claims("audit_logging", r"\bretain(?:ed)?\s+(?:for\s+)?(?:a\s+minimum\s+of\s+)?\d+\s+(?:months?|years?)\b|\b\d+[- ](?:month|year)\s+retention\b"):
            out.append(
                "The 'audit_logging' answer commits to a specific log retention period. "
                "Confirm against policy or replace with [Verify: retention period]."
            )
        # Breach-window mismatch in the answer itself (e.g. AI answered "36 months"
        # when the question asks about 24 months).
        if claims("breach_history", r"\bpast\s+(?:3[6-9]|[4-9]\d|\d{3,})\s+months?\b|\b(?:3|4|5)\s+years?\b"):
            out.append(
                "The 'breach_history' answer references a window longer than the 24 months the "
                "question asks about. Re-answer for the 24-month window only."
            )

        # Cloud providers — flag AWS / Azure / GCP / OCI when intake silent
        if not intake.get("primary_cloud"):
            for key in ("data_residency", "security_measures", "encryption_standards"):
                if claims(key, r"\baws\b|\bamazon web services\b|\bazure\b|\bmicrosoft azure\b|\bgcp\b|\bgoogle cloud\b|\boracle cloud\b|\boci\b"):
                    out.append(
                        f"The '{key}' answer names a specific cloud provider (AWS/Azure/GCP/etc.) but "
                        "the intake did not declare one. Confirm the actual provider or use a placeholder."
                    )
                    break

        # PDPC breach-notification window — must be 3 calendar days (PDPA §26D),
        # NOT "1 hour" / "immediately" / "24 hours" / "to the government".
        ir_ans = lower("incident_response")
        if ir_ans:
            if re.search(r"\bwithin\s+1\s+hour\b|\bwithin\s+24\s+hours?\b|\bimmediately\s+(?:notify|inform)\b", ir_ans):
                out.append(
                    "The 'incident_response' answer commits to a breach-notification window that does "
                    "not match PDPA §26D (3 calendar days for notifiable breaches). Correct the window."
                )
            if re.search(r"\bnotify(?:ing)?\s+the\s+government\b|\binform(?:ing)?\s+the\s+government\b", ir_ans):
                out.append(
                    "The 'incident_response' answer says breaches are notified to 'the government'. "
                    "Under PDPA §26D, notification is to the PDPC (Personal Data Protection Commission), "
                    "not 'the government'. Correct the recipient."
                )

        return out

    # ── Step 2.5: blockchain anchor ───────────────────────────────────────────

    def _compute_evidence_hash(self, company_name: str, qa_answers: Dict[str, str]) -> str:
        """Deterministic SHA-256 over the kit's actual content.

        This is the value we both ANCHOR on-chain and PRINT in the PDF, so a
        procurer can take the hash off the document and verify it against the
        EvidenceAnchorV3 contract. Previously the PDF printed the raw UUID
        report_id (not a hash at all) while the chain anchored sha256(report_id)
        — two different strings, so the "evidence hash" couldn't be verified.

        Binding the hash to report_id + company + the sorted Q&A means it proves
        the specific answers existed at anchor time, not merely that an id was
        minted. Sorting keys keeps it stable across regeneration of identical
        content (idempotent re-anchor skip).
        """
        import hashlib
        import json as _json
        payload = {
            "report_id": self.report_id,
            "company": company_name or "",
            "qa": {k: qa_answers[k] for k in sorted(qa_answers.keys())},
        }
        canonical = _json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
        return hashlib.sha256(canonical.encode("utf-8")).hexdigest()

    # ── Intake-fill + quality gate ────────────────────────────────────────────

    # Placeholder markers the AI/template emit for facts the buyer must supply.
    _PLACEHOLDER_RE = re.compile(r"\[\s*(?:FILL IN|Verify:)[^\]]*\]|_{2,}\s*\[\s*FILL IN\s*\]\s*_{2,}", re.I)

    def _apply_intake_substitutions(self, qa_answers: Dict[str, str], intake: dict | None) -> Dict[str, str]:
        """Replace `[Verify: …]` placeholders with facts the buyer gave in intake.

        Only fills placeholders whose underlying fact the buyer actually supplied
        — anything genuinely unknown stays a placeholder for them to complete.
        Matching is keyed off the placeholder's own wording so we don't overwrite
        an unrelated field.
        """
        if not intake or not isinstance(qa_answers, dict):
            return qa_answers

        def _v(key: str) -> str:
            val = (intake.get(key) or "").strip() if isinstance(intake.get(key), str) else intake.get(key)
            return val if val and val not in ("unknown", "") else ""

        subs: list[tuple[re.Pattern, str]] = []
        iso_num, iso_exp = _v("iso_cert_number"), _v("iso_cert_expiry")
        if iso_num:
            iso_text = f"ISO/IEC 27001 certificate {iso_num}"
            if iso_exp:
                iso_text += f" (valid until {iso_exp})"
            subs.append((re.compile(r"\[\s*Verify:[^\]]*ISO[^\]]*\]", re.I), iso_text))
        if _v("bcp_last_tested"):
            subs.append((
                re.compile(r"\[\s*Verify:[^\]]*(?:BCP|business continuity|DR plan|last test(?:ed)?)[^\]]*\]", re.I),
                f"last tested {_v('bcp_last_tested')}",
            ))
        # New-hire training window — must precede the generic training-cadence
        # pattern below so "[Verify: new-hire training window]" is not consumed by
        # the cadence rule.
        if _v("training_newhire_window"):
            subs.append((
                re.compile(r"\[\s*Verify:[^\]]*(?:new[- ]?hire|new[- ]?joiner|onboarding)[^\]]*\]", re.I),
                f"new-hire training within {_v('training_newhire_window')}",
            ))
        if _v("training_frequency"):
            # Anchor cadence/frequency right after "Verify:" (optionally prefixed
            # with "training") so this does NOT swallow other cadence placeholders
            # like "[Verify: scan cadence]" or "[Verify: access review cadence]".
            subs.append((
                re.compile(r"\[\s*Verify:\s*(?:training\s+)?(?:cadence|frequency|awareness training)[^\]]*\]", re.I),
                str(_v("training_frequency")),
            ))
        if _v("key_processors"):
            subs.append((
                re.compile(r"\[\s*Verify:[^\]]*(?:sub-?processors?|processors?)[^\]]*\]", re.I),
                str(_v("key_processors")),
            ))

        # ── Remaining categories (audit fix: full placeholder coverage) ──────
        # Each fills the [Verify: …] marker the AI emits for that fact so the
        # hard gate is satisfiable once the buyer supplies it in intake.
        if _v("soc2_status"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*SOC[\s-]?2[^\]]*\]", re.I),
                         f"SOC 2 status: {_v('soc2_status')}"))
        _dpo_name, _dpo_email, _dpo_reg = _v("dpo_name"), _v("dpo_email"), _v("dpo_pdpc_reg")
        if _dpo_reg:
            subs.append((re.compile(r"\[\s*Verify:[^\]]*PDPC[^\]]*\]", re.I),
                         f"PDPC registration {_dpo_reg}"))
        if _dpo_name or _dpo_email:
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:DPO|data protection officer)[^\]]*\]", re.I),
                         "; ".join(p for p in [_dpo_name, _dpo_email] if p)))
        if _v("bcp_rto"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*RTO[^\]]*\]", re.I), f"RTO {_v('bcp_rto')}"))
        if _v("bcp_rpo"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*RPO[^\]]*\]", re.I), f"RPO {_v('bcp_rpo')}"))
        if _v("access_review_cadence"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:access review|privileged access|review cadence)[^\]]*\]", re.I),
                         f"privileged access reviewed {_v('access_review_cadence')}"))
        if _v("mfa_privileged"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:MFA|multi[- ]?factor)[^\]]*\]", re.I),
                         str(_v("mfa_privileged"))))
        if _v("patch_sla"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:patch|remediation)[^\]]*\]", re.I),
                         f"critical patches within {_v('patch_sla')}"))
        if _v("scan_cadence"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:scan|vulnerability)[^\]]*\]", re.I),
                         f"vulnerability scans {_v('scan_cadence')}"))
        # Encryption — specific (at rest / in transit / key mgmt) before generic.
        if _v("encryption_at_rest"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:encryption at rest|at[- ]rest)[^\]]*\]", re.I),
                         f"{_v('encryption_at_rest')} at rest"))
        if _v("encryption_in_transit"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:encryption in transit|in[- ]transit)[^\]]*\]", re.I),
                         f"{_v('encryption_in_transit')} in transit"))
        if _v("key_management"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:key management|key[- ]mgmt)[^\]]*\]", re.I),
                         str(_v("key_management"))))
        if _v("encryption_at_rest") or _v("encryption_in_transit"):
            _enc = " / ".join(p for p in [
                (f"{_v('encryption_at_rest')} at rest" if _v("encryption_at_rest") else ""),
                (f"{_v('encryption_in_transit')} in transit" if _v("encryption_in_transit") else ""),
            ] if p)
            subs.append((re.compile(r"\[\s*Verify:[^\]]*encryption[^\]]*\]", re.I), _enc))
        if _v("log_retention"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:retention|audit log|log retention)[^\]]*\]", re.I),
                         f"logs retained {_v('log_retention')}"))
        if _v("log_monitoring"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:monitoring|anomaly)[^\]]*\]", re.I),
                         str(_v("log_monitoring"))))
        if _v("incident_notification_window"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:notification window|internal notification|incident)[^\]]*\]", re.I),
                         f"internal notification within {_v('incident_notification_window')}"))
        if _v("cross_border_mechanism"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:cross[- ]border|data cent(?:re|er)|hosting region|transfer mechanism)[^\]]*\]", re.I),
                         str(_v("cross_border_mechanism"))))
        if _v("subcontracting"):
            subs.append((re.compile(r"\[\s*Verify:[^\]]*(?:subcontract|offshor)[^\]]*\]", re.I),
                         str(_v("subcontracting"))))

        if not subs:
            return qa_answers

        filled = 0
        out: Dict[str, str] = {}
        for k, v in qa_answers.items():
            if isinstance(v, str):
                for pattern, repl in subs:
                    v, n = pattern.subn(repl, v)
                    filled += n
            out[k] = v
        if filled:
            logger.info(f"[RFP] Filled {filled} intake-backed placeholder(s) in qa_answers")
        return out

    def _count_residual_placeholders(self, qa_answers: Dict[str, str]) -> int:
        """Count `[FILL IN]` / `[Verify: …]` markers still present after fill."""
        if not isinstance(qa_answers, dict):
            return 0
        return sum(
            len(self._PLACEHOLDER_RE.findall(v))
            for v in qa_answers.values()
            if isinstance(v, str)
        )

    # Instructional guidance per placeholder topic — tells the buyer exactly
    # WHERE to find the fact and in what FORMAT to enter it, instead of leaving
    # them with a bare "[Verify: encryption standard]" marker (forensic-audit
    # finding: placeholders gave no guidance). Matched against the marker text;
    # first hit wins, so order most-specific → generic.
    _PLACEHOLDER_GUIDANCE: list[tuple[re.Pattern, str]] = [
        (re.compile(r"ISO", re.I),
         "Check your ISO certificate. Format: ISO/IEC 27001:2022 — Certificate No. [number] — valid until [date]."),
        (re.compile(r"SOC[\s-]?2", re.I),
         "Check your SOC 2 report cover page. Format: SOC 2 Type II — report period [dates]."),
        (re.compile(r"PDPC", re.I),
         "Your DPO's PDPC registration reference, if filed with the PDPC."),
        (re.compile(r"DPO|data protection officer", re.I),
         "Your appointed Data Protection Officer. Format: [name] — [email]."),
        (re.compile(r"encryption|at[- ]rest|in[- ]transit|key management|key[- ]mgmt", re.I),
         "Check AWS Console > Security Hub (or your IT manager's encryption policy). Format: AES-256 at rest / TLS 1.2+ in transit; keys in [AWS KMS / customer-managed]."),
        (re.compile(r"RTO|RPO|BCP|business continuity|DR plan|last test", re.I),
         "From your BCP/DR test report. Format: last tested [date]; RTO [hours]; RPO [hours]."),
        (re.compile(r"patch|remediation", re.I),
         "From your patch-management policy. Format: critical patches applied within [N] days."),
        (re.compile(r"scan|vulnerabilit", re.I),
         "From your vulnerability-management policy. Format: vulnerability scans [monthly/quarterly]."),
        (re.compile(r"new[- ]?hire|new[- ]?joiner|onboarding", re.I),
         "From your training records. Format: new-hire security training within [N] days of joining."),
        (re.compile(r"training|awareness", re.I),
         "From your training records. Format: staff security-awareness training [annually/quarterly]."),
        (re.compile(r"access review|privileged access|review cadence|MFA|multi[- ]?factor", re.I),
         "From your IAM/access-review log. Format: privileged access reviewed [quarterly]; MFA enforced on privileged accounts [yes/no]."),
        (re.compile(r"retention|audit log|log retention|monitoring|anomaly", re.I),
         "From your logging policy. Format: logs retained [N months]; monitored via [SIEM/tool]."),
        (re.compile(r"notification window|internal notification|incident", re.I),
         "From your incident-response plan. Format: internal breach notification within [N hours]."),
        (re.compile(r"cross[- ]border|data cent(?:re|er)|hosting region|transfer mechanism", re.I),
         "From your cloud config. Format: data hosted in [region]; cross-border transfers under [SCCs / other mechanism]."),
        (re.compile(r"sub-?processor|processor|subcontract|offshor", re.I),
         "From your vendor list. Format: key sub-processors [names]; subcontracting/offshoring [yes/no + where]."),
    ]

    def _guidance_for_marker(self, marker: str) -> str | None:
        """Where-to-find + expected-format hint for a placeholder marker, if any."""
        for pattern, hint in self._PLACEHOLDER_GUIDANCE:
            if pattern.search(marker):
                return hint
        return None

    def _residual_placeholder_details(self, qa_answers: Dict[str, str]) -> list[str]:
        """Return the distinct surviving placeholder markers, each enriched with
        instructional guidance, in order.

        The marker names the fact the buyer must supply (e.g.
        "[Verify: ISO 27001 cert number and expiry]"); we append a where-to-find
        + format hint so the buyer knows exactly how to complete it when delivery
        is blocked and they are routed back to the intake.
        """
        if not isinstance(qa_answers, dict):
            return []
        seen: list[str] = []
        seen_markers: set[str] = set()
        for v in qa_answers.values():
            if not isinstance(v, str):
                continue
            for marker in self._PLACEHOLDER_RE.findall(v):
                label = " ".join(marker.split()).strip()
                if not label or label in seen_markers:
                    continue
                seen_markers.add(label)
                hint = self._guidance_for_marker(label)
                seen.append(f"{label} — {hint}" if hint else label)
        return seen

    async def _anchor_to_blockchain(self) -> Optional[str]:
        try:
            from app.services.blockchain import BlockchainService
            from app.core.config import settings
            # self.evidence_hash is a 64-char SHA-256 hex computed from the kit
            # content (see _compute_evidence_hash) — valid bytes32 input AND the
            # exact value printed in the PDF so the anchor is independently
            # verifiable. Defensive fallback to sha256(report_id) if a caller
            # anchors before the hash is set.
            import hashlib
            evidence_hash = getattr(self, "evidence_hash", None) or hashlib.sha256(
                self.report_id.encode()
            ).hexdigest()
            self.evidence_hash = evidence_hash
            # Admin test-checkout runs (admin-sim-* session) mock the anchor so QA
            # never spends real gas from the shared wallet.
            from app.core.demo_flags import is_demo_anchor
            blockchain = BlockchainService()
            tx = await blockchain.anchor_evidence(
                evidence_hash,
                metadata=f"rfp_express:vendor:{self.vendor_id}",
                demo=is_demo_anchor(session_id=self.session_id),
            )
            logger.info(f"RFP Express anchored on {settings.active_polygon_network_name}: {tx}")
            return tx
        except Exception as e:
            logger.warning(f"Blockchain anchor failed for RFP Express (non-blocking): {e}")
            self.warnings.append(f"Blockchain anchor skipped: {e}")
            return None

    # ── Step 3: build PDF ─────────────────────────────────────────────────────

    def _build_pdf(
        self,
        company_name: str,
        vendor_url: str,
        qa_answers: Dict[str, str],
        ctx: Dict,
        tx_hash: Optional[str] = None,
        product_type: str = "rfp_express",
        acra_live: Dict | None = None,
        pdpc_result: Dict | None = None,
        discrepancies: list | None = None,
        intake: dict | None = None,
        verification_map: Dict[str, Dict[str, Any]] | None = None,
        coverage_summary: str | None = None,
        gebiz_history: Dict | None = None,
        dns_security: Dict | None = None,
        onemap_location: Dict | None = None,
    ) -> bytes:
        try:
            from app.services.pdf_service import PDFService
            from app.core.config import settings
            pdf = PDFService()
            verify_base = settings.VERIFY_BASE_URL.rstrip("/")

            # Pre-flight review checklist — one item per answer the buyer must
            # tick off. Rendered as a real checkbox table by PDFService rather
            # than a run-on paragraph. Forces manual acknowledgement of each
            # answer before submission instead of paste-and-pray.
            checklist_items = [
                f"{self._q_label(k)} — reviewed, accurate, no [FILL IN] left."
                for k in qa_answers.keys()
            ]
            checklist_confirmations = [
                "All [FILL IN] placeholders have been replaced with real, accurate values.",
                "I confirm the company has the capability to substantiate every claim above.",
                "This document is being submitted alongside the proposal, pricing, and team sections.",
            ]

            # Verification label per answer — mirrors the web result page.
            # Buyers reading the PDF see the same source attribution.
            _SOURCE_LABEL = {
                "intake":           "From your intake",
                "website":          "Verified on your website",
                "intake+website":   "Intake + website verified",
                "intake+external":  "Intake + public records",
                "acra":             "ACRA verified",
                "ssl":              "SSL Labs verified",
                "gebiz":            "GeBIZ supplier verified",
                "pdpc":             "PDPC register checked",
                "external":         "External evidence verified",
                "ai_drafted":       "AI draft — review before submission",
            }
            verification_map = verification_map or {}

            def _verif_text(key: str) -> str:
                """Bracket-free verification caption for the structured renderer."""
                vinfo = verification_map.get(key) or {"source": "ai_drafted", "evidence": []}
                src = vinfo.get("source") or "ai_drafted"
                label = _SOURCE_LABEL.get(src, src)
                ev = vinfo.get("evidence") or []
                ev_part = (" · Evidence: " + " · ".join(ev[:3])) if ev else ""
                return f"{label}{ev_part}"

            qa_list = [
                {
                    "question": self._q_label(k),
                    "answer": v,
                    "verification": _verif_text(k),
                    "ai_drafted": (verification_map.get(k) or {}).get("source", "ai_drafted") == "ai_drafted",
                }
                for k, v in qa_answers.items()
            ]

            is_complete = product_type == "rfp_complete"
            framework_label = "RFP Kit Complete Evidence Pack" if is_complete else "RFP Kit Express Evidence Certificate"

            intake = intake or {}

            # ── Structured evidence detail rows (symbol-free; PDFService renders
            # these in a meta table). Blockchain anchor is rendered separately by
            # PDFService._blockchain_block, so it is not duplicated here.
            details: list[tuple[str, str]] = [
                ("Vendor URL", vendor_url),
                ("Sector", ctx.get("sector") or "General"),
                ("UEN (Business Reg. No.)", ctx.get("uen") or "Not provided"),
            ]

            # Tender personalization — when the buyer told us which tender they
            # are bidding on, attribute the kit to it up front.
            _tender_line = self._tender_attribution(intake)
            if _tender_line:
                details.insert(0, ("Prepared in response to", _tender_line))

            # ACRA verification
            if acra_live and acra_live.get("found"):
                entity_type = acra_live.get("entity_type", "")
                acra_val = ("LIVE" if acra_live.get("live")
                            else (acra_live.get("entity_status") or "Inactive"))
                details.append(("ACRA Status", f"{acra_val}{f' ({entity_type})' if entity_type else ''}"))
            else:
                details.append(("ACRA Status", "Not verified"))

            if coverage_summary:
                details.append(("Coverage", coverage_summary))

            # GeBIZ supplier
            if ctx.get("gebiz_supplier"):
                count = ctx.get("gebiz_contracts_count") or 0
                details.append((
                    "GeBIZ Supplier",
                    f"Registered — {count} prior government contract(s)" if count
                    else "Registered",
                ))

            # DPO contact (audit fix B)
            dpo_parts = [x for x in [intake.get("dpo_name") or "", intake.get("dpo_email") or ""] if x]
            if dpo_parts:
                details.append(("DPO Contact", ", ".join(dpo_parts)))

            # Privacy policy URL (audit fix C/F)
            if ctx.get("privacy_policy_url"):
                details.append(("Privacy Policy", ctx["privacy_policy_url"]))

            # ISO certification
            if intake.get("iso_cert_number"):
                expiry = intake.get("iso_cert_expiry", "")
                details.append((
                    "ISO 27001 Cert",
                    f"{intake['iso_cert_number']}"
                    + (f" (expires {expiry})" if expiry else "")
                    + " — verify at bsigroup.com/en-SG/validate-bsi-issued-certificates/",
                ))
            elif intake.get("iso_status") and intake["iso_status"].lower() not in ("no", "none", "pursuing", ""):
                details.append((
                    "ISO Status",
                    f"{intake['iso_status']} — certificate number not provided; "
                    "buyers should request the cert for independent verification",
                ))

            # PDPC enforcement
            if pdpc_result and pdpc_result.get("found"):
                details.append(("PDPC Enforcement", "Previous enforcement action found — see warnings"))
                
            # Additional Enrichments
            if gebiz_history and gebiz_history.get("total_awards", 0) > 0:
                details.append(("GeBIZ Records", f"{gebiz_history['total_awards']} historical contract awards"))
                
            if dns_security and dns_security.get("checked"):
                if dns_security.get("dmarc_record"):
                    details.append(("Email Security", f"DMARC Enforced (p={dns_security.get('dmarc_policy')})"))
                elif dns_security.get("spf_record"):
                    details.append(("Email Security", "SPF configured, DMARC missing"))
                    
            if onemap_location and onemap_location.get("found"):
                details.append(("Verified Address", onemap_location.get("planning_area")))

            # Scope notice — kit covers PDPA + security only; buyer must add the
            # rest of the bid (proposal, pricing, team) themselves.
            scope_intro = (
                "This document covers PDPA and information-security compliance answers only. "
                "It is NOT a complete bid response. Before submitting to GeBIZ or any other "
                "procurement portal, you must add your own:"
            )
            scope_bullets = [
                "Technical proposal",
                "Pricing and commercial terms",
                "Delivery timeline and milestones",
                "Team and key personnel",
                "Tender-specific requirements (references, case studies, certifications attached)",
            ]
            scope_closing = (
                "Review every answer below and replace any [FILL IN] placeholders before submission. "
                "Booppa does not warrant the suitability of these answers for any particular tender; "
                "the buyer remains responsible for the accuracy of submitted statements."
            )

            report_data = {
                "company_name": company_name,
                "report_id":    self.report_id,
                "created_at":   datetime.now(timezone.utc).isoformat(),
                "framework":    framework_label,
                "product_type": product_type,
                "status":       "Completed",
                "summary": (
                    f"This certificate confirms that {company_name} has completed the "
                    f"BOOPPA {'RFP Kit Complete' if is_complete else 'RFP Kit Express'} process, "
                    f"generating blockchain-anchored evidence for procurement submission."
                    + (f" Prepared in response to {_tender_line}." if _tender_line else "")
                ),
                # Structured payload consumed by PDFService._rfp_kit_story.
                "rfp_kit": {
                    "scope_intro": scope_intro,
                    "scope_bullets": scope_bullets,
                    "scope_closing": scope_closing,
                    "details": details,
                    "checklist": checklist_items,
                    "checklist_confirmations": checklist_confirmations,
                    "qa": qa_list,
                    "template_used": bool(self.used_template),
                    "discrepancies": list(discrepancies or []),
                },
                # The evidence hash shown to the buyer is the SAME SHA-256 we
                # anchored on-chain (not the raw UUID report_id) — so it is
                # actually verifiable against the EvidenceAnchorV3 contract.
                # report_id remains available above as the human-facing Report ID.
                "audit_hash": getattr(self, "evidence_hash", None) or self.report_id,
                "verify_url": f"{verify_base}/verify/{self.report_id}",
                "tx_hash": tx_hash,
            }
            return pdf.generate_pdf(report_data)
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            self.errors.append(f"PDF generation error: {e}")
            raise

    def _q_label(self, key: str) -> str:
        return QUESTION_LABELS.get(key, key.replace("_", " ").title())

    @staticmethod
    def _tender_attribution(intake: dict | None) -> str:
        """Human-readable 'this kit was prepared for tender X by agency Y' line.

        Personalizes the RFP kit to the specific tender the buyer is bidding on,
        drawn from the tender fields the intake captured (extracted from an
        uploaded tender PDF or typed by the buyer). Returns "" when no tender
        identifiers are on file so the kit stays generic rather than inventing one.
        """
        intake = intake or {}
        ref = (intake.get("tender_ref") or "").strip()
        agency = (intake.get("tender_agency") or "").strip()
        title = (intake.get("tender_title") or "").strip()
        parts: list[str] = []
        if ref:
            parts.append(f"Tender {ref}")
        if title:
            parts.append(f'"{title}"' if not ref else title)
        if agency:
            parts.append(f"issued by {agency}")
        return " — ".join(parts) if parts else ""

    def _compute_verification(
        self,
        intake: dict,
        vendor_ctx: dict,
        website_signals: Dict[str, Any],
        ssl_result: Dict | None = None,
        domain_rep: Dict | None = None,
        acra_live: Dict | None = None,
    ) -> Dict[str, Dict[str, Any]]:
        """Compute {key: {source, evidence}} for each RFP question.

        Source values used by the frontend to pick a badge label:
          • 'intake'          → "From your intake"
          • 'website'         → "Verified on your website"
          • 'acra' / 'gebiz'  → "ACRA verified" / "GeBIZ supplier"
          • 'ssl'             → "SSL Labs grade X"
          • 'intake+website'  → highest trust — buyer + their published site agree
          • 'intake+external' → buyer claim + ACRA/GeBIZ/etc.
          • 'ai_drafted'      → no backing evidence (amber "AI draft" badge)

        Each value carries an `evidence` list — human-readable strings the
        result page surfaces on hover and the PDF prints under the answer.
        Priority chain: if a question has multiple sources, the strongest
        wins; intake+website is the highest combo.
        """
        intake = intake or {}
        ws = website_signals or {}
        ctx = vendor_ctx or {}
        _empty = (None, "unknown", "")

        out: Dict[str, Dict[str, Any]] = {}

        def attach(key: str, src: str, evidence: str) -> None:
            cur = out.get(key)
            if not cur:
                out[key] = {"source": src, "evidence": [evidence]}
                return
            cur["evidence"].append(evidence)
            # Merge sources: prefer combined labels when intake meets an
            # external/website source.
            existing = cur["source"]
            if src == existing:
                return
            paired = {existing, src}
            if "intake" in paired and "website" in paired:
                cur["source"] = "intake+website"
            elif "intake" in paired and src in {"acra", "ssl", "gebiz", "pdpc", "external"}:
                cur["source"] = "intake+external"
            elif existing in {"acra", "ssl", "gebiz", "pdpc"} and src in {"acra", "ssl", "gebiz", "pdpc"}:
                cur["source"] = "external"
            # else keep the stronger of (intake > website > external)

        # ── data_policy ─────────────────────────────────────────────────────
        if ctx.get("uen") or intake.get("uen"):
            attach("data_policy", "intake", f"Company UEN provided ({ctx.get('uen') or intake.get('uen')})")
        if acra_live and acra_live.get("found"):
            attach("data_policy", "acra", f"ACRA verified · {acra_live.get('registered_name') or 'live entity'}")
        if ws.get("pdpa_mentioned"):
            attach("data_policy", "website", "PDPA referenced on your website")
        if ctx.get("privacy_policy_url"):
            attach("data_policy", "website", f"Privacy policy published at {ctx['privacy_policy_url']}")

        # ── dpo_appointed ──────────────────────────────────────────────────
        if intake.get("dpo_appointed") not in _empty:
            attach("dpo_appointed", "intake", f"Intake: dpo_appointed={intake['dpo_appointed']}")
        if intake.get("dpo_name") not in _empty or intake.get("dpo_email") not in _empty:
            attach("dpo_appointed", "intake",
                   f"DPO contact supplied: {intake.get('dpo_name') or ''} {intake.get('dpo_email') or ''}".strip())
        if ws.get("dpo_mentioned"):
            attach("dpo_appointed", "website", "DPO referenced on your website")

        # ── security_measures ──────────────────────────────────────────────
        if ssl_result and ssl_result.get("grade"):
            attach("security_measures", "ssl", f"SSL Labs grade {ssl_result['grade']}")
        if ws.get("encryption_generic") or ws.get("aes_mentioned") or ws.get("tls_mentioned"):
            terms = []
            if ws.get("aes_mentioned"): terms.append("AES")
            if ws.get("tls_mentioned"): terms.append("TLS")
            if not terms: terms.append("encryption")
            attach("security_measures", "website", f"Security language on your website: {', '.join(terms)}")
        if ws.get("iso_27001_mentioned"):
            attach("security_measures", "website", "ISO 27001 referenced on your website")

        # ── breach_history ─────────────────────────────────────────────────
        if intake.get("breach_history") not in _empty:
            attach("breach_history", "intake", f"Intake: breach_history={intake['breach_history']}")
        if domain_rep and domain_rep.get("checked"):
            flagged = domain_rep.get("flagged")
            attach("breach_history", "external",
                   "VirusTotal: domain clean" if not flagged else f"VirusTotal: flagged by {domain_rep.get('malicious_votes', 0)} vendor(s)")
        # No PDPC enforcement check is positive evidence for a no-breach claim
        # only when we have a clean result, not an absence of data.

        # ── third_party ────────────────────────────────────────────────────
        if intake.get("key_processors") not in _empty:
            attach("third_party", "intake", f"Key processors listed: {intake['key_processors']}")
        if ws.get("subprocessors_mentioned"):
            attach("third_party", "website", "Sub-processors referenced on your website / policy")
        if ws.get("dpa_mentioned"):
            attach("third_party", "website", "Data Processing Agreement referenced on your website")

        # ── iso_certifications ─────────────────────────────────────────────
        if intake.get("iso_status") not in _empty:
            attach("iso_certifications", "intake", f"Intake: iso_status={intake['iso_status']}")
        if intake.get("iso_cert_number") not in _empty:
            attach("iso_certifications", "intake",
                   f"ISO cert {intake['iso_cert_number']}" +
                   (f" (exp {intake['iso_cert_expiry']})" if intake.get("iso_cert_expiry") else ""))
        if ws.get("iso_27001_mentioned"):
            year = ws.get("iso_27001_year")
            attach("iso_certifications", "website",
                   f"ISO 27001{':' + year if year else ''} referenced on your website")
        if ws.get("soc_2_mentioned"):
            attach("iso_certifications", "website", "SOC 2 referenced on your website")
        if ws.get("iso_27701_mentioned"):
            attach("iso_certifications", "website", "ISO 27701 referenced on your website")
        if intake.get("soc2_status") not in _empty:
            attach("iso_certifications", "intake", f"Intake: soc2_status={intake['soc2_status']}")

        # ── dpo PDPC registration ──────────────────────────────────────────
        if intake.get("dpo_pdpc_reg") not in _empty:
            attach("dpo_appointed", "intake", f"DPO PDPC registration: {intake['dpo_pdpc_reg']}")

        # ── business_continuity ────────────────────────────────────────────
        if intake.get("bcp_last_tested") not in _empty:
            attach("business_continuity", "intake", f"BCP last tested: {intake['bcp_last_tested']}")
        if intake.get("bcp_rto") not in _empty:
            attach("business_continuity", "intake", f"RTO target: {intake['bcp_rto']}")
        if intake.get("bcp_rpo") not in _empty:
            attach("business_continuity", "intake", f"RPO target: {intake['bcp_rpo']}")

        # ── staff_training ─────────────────────────────────────────────────
        if intake.get("training_frequency") not in _empty:
            attach("staff_training", "intake", f"Training frequency: {intake['training_frequency']}")
        if intake.get("training_newhire_window") not in _empty:
            attach("staff_training", "intake", f"New-hire training window: {intake['training_newhire_window']}")

        # ── access_controls — intake-only when supplied
        if intake.get("access_review_cadence") not in _empty:
            attach("access_controls", "intake", f"Privileged access review cadence: {intake['access_review_cadence']}")
        if intake.get("mfa_privileged") not in _empty:
            attach("access_controls", "intake", f"MFA on privileged accounts: {intake['mfa_privileged']}")

        # ── vulnerability_mgmt — intake-only when supplied
        if intake.get("patch_sla") not in _empty:
            attach("vulnerability_mgmt", "intake", f"Critical patch SLA: {intake['patch_sla']}")
        if intake.get("scan_cadence") not in _empty:
            attach("vulnerability_mgmt", "intake", f"Vulnerability scan cadence: {intake['scan_cadence']}")

        # ── encryption_standards ───────────────────────────────────────────
        if ssl_result and ssl_result.get("grade"):
            attach("encryption_standards", "ssl", f"SSL Labs grade {ssl_result['grade']}")
        if ws.get("aes_mentioned") or ws.get("tls_mentioned"):
            terms = []
            if ws.get("aes_mentioned"): terms.append("AES")
            if ws.get("tls_mentioned"): terms.append("TLS")
            attach("encryption_standards", "website",
                   f"{', '.join(terms)} mentioned on your website")
        if intake.get("encryption_at_rest") not in _empty:
            attach("encryption_standards", "intake", f"Encryption at rest: {intake['encryption_at_rest']}")
        if intake.get("encryption_in_transit") not in _empty:
            attach("encryption_standards", "intake", f"Encryption in transit: {intake['encryption_in_transit']}")
        if intake.get("key_management") not in _empty:
            attach("encryption_standards", "intake", f"Key management: {intake['key_management']}")

        # ── audit_logging — intake-only when supplied
        if intake.get("log_retention") not in _empty:
            attach("audit_logging", "intake", f"Audit log retention: {intake['log_retention']}")
        if intake.get("log_monitoring") not in _empty:
            attach("audit_logging", "intake", f"Log monitoring: {intake['log_monitoring']}")

        # ── incident_response ──────────────────────────────────────────────
        if ws.get("breach_policy_mentioned"):
            attach("incident_response", "website", "Incident response / breach policy referenced on your website")
        if intake.get("incident_notification_window") not in _empty:
            attach("incident_response", "intake", f"Internal notification window: {intake['incident_notification_window']}")

        # ── data_residency ─────────────────────────────────────────────────
        if intake.get("data_hosting") not in _empty:
            attach("data_residency", "intake", f"Intake hosting: {intake['data_hosting']}")
        if intake.get("primary_cloud") not in _empty:
            attach("data_residency", "intake", f"Primary cloud: {intake['primary_cloud']}")
        if intake.get("cross_border_mechanism") not in _empty:
            attach("data_residency", "intake", f"Cross-border transfer mechanism: {intake['cross_border_mechanism']}")
        if ws.get("singapore_residency_mentioned"):
            attach("data_residency", "website", "Singapore data residency referenced on your website")
        clouds = []
        if ws.get("aws_mentioned"): clouds.append("AWS")
        if ws.get("azure_mentioned"): clouds.append("Azure")
        if ws.get("gcp_mentioned"): clouds.append("GCP")
        if clouds:
            attach("data_residency", "website", f"Cloud provider(s) named on your website: {', '.join(clouds)}")
        # Hosting header inference is also evidence
        if ctx.get("inferred_hosting_provider"):
            region = f" ({ctx['inferred_hosting_region']})" if ctx.get("inferred_hosting_region") else ""
            attach("data_residency", "external",
                   f"HTTP headers infer hosting: {ctx['inferred_hosting_provider']}{region}")

        # ── subcontracting ─────────────────────────────────────────────────
        if intake.get("key_processors") not in _empty:
            attach("subcontracting", "intake", f"Processors listed: {intake['key_processors']}")
        if intake.get("subcontracting") not in _empty:
            attach("subcontracting", "intake", f"Subcontracting: {intake['subcontracting']}")
        if acra_live and acra_live.get("found"):
            attach("subcontracting", "acra", "Entity verified on ACRA")
        if ctx.get("gebiz_supplier"):
            count = ctx.get("gebiz_contracts_count") or 0
            attach("subcontracting", "gebiz",
                   f"GeBIZ registered supplier · {count} prior government contract(s)" if count else "GeBIZ registered supplier")
        if ws.get("subprocessors_mentioned"):
            attach("subcontracting", "website", "Sub-processors policy referenced on your website")

        return out

    def _fact_backed_keys(self, intake: dict, ctx: dict,
                          ssl_result: Dict | None = None,
                          domain_rep: Dict | None = None,
                          acra_live: Dict | None = None) -> set:
        """Return the set of question keys that have a real fact grounding them."""
        _empty = (None, "unknown", "")
        backed = set()
        # DPO — backed if status provided or name/email supplied
        if intake.get("dpo_appointed") not in _empty:
            backed.add("dpo_appointed")
        if intake.get("dpo_name") not in _empty or intake.get("dpo_email") not in _empty:
            backed.add("dpo_appointed")  # named DPO is stronger backing
        # ISO — backed if status or cert number provided
        if intake.get("iso_status") not in _empty:
            backed.add("iso_certifications")
        if intake.get("iso_cert_number") not in _empty:
            backed.add("iso_certifications")
        # Data residency — backed if hosting, cloud provider, or region provided
        if intake.get("data_hosting") not in _empty:
            backed.add("data_residency")
        if intake.get("primary_cloud") not in _empty or intake.get("cloud_region") not in _empty:
            backed.add("data_residency")
        # Breach history — backed by declaration or external check
        if intake.get("breach_history") not in _empty or (domain_rep and domain_rep.get("checked")):
            backed.add("breach_history")
        # BCP / staff training — backed if vendor provided specific dates/frequency
        if any(intake.get(k) not in _empty for k in ("bcp_last_tested", "bcp_rto", "bcp_rpo")):
            backed.add("business_continuity")
        if intake.get("training_frequency") not in _empty or intake.get("training_newhire_window") not in _empty:
            backed.add("staff_training")
        # SOC 2 strengthens iso_certifications
        if intake.get("soc2_status") not in _empty:
            backed.add("iso_certifications")
        # Access controls / vuln mgmt / encryption / audit logging / incident
        # response — intake-only categories; backed once the buyer supplies them.
        if any(intake.get(k) not in _empty for k in ("access_review_cadence", "mfa_privileged")):
            backed.add("access_controls")
        if any(intake.get(k) not in _empty for k in ("patch_sla", "scan_cadence")):
            backed.add("vulnerability_mgmt")
        if any(intake.get(k) not in _empty for k in ("encryption_at_rest", "encryption_in_transit", "key_management")):
            backed.add("encryption_standards")
        if any(intake.get(k) not in _empty for k in ("log_retention", "log_monitoring")):
            backed.add("audit_logging")
        if intake.get("incident_notification_window") not in _empty:
            backed.add("incident_response")
        if intake.get("cross_border_mechanism") not in _empty:
            backed.add("data_residency")
        # Third-party — backed if key processors listed
        if intake.get("key_processors") not in _empty:
            backed.add("third_party")
        if intake.get("subcontracting") not in _empty:
            backed.add("subcontracting")
        # Data policy — backed by ACRA/UEN (entity verification)
        if ctx.get("uen") or intake.get("uen"):
            backed.add("data_policy")
        if ctx.get("acra_name") or (acra_live and acra_live.get("found")):
            backed.update({"data_policy", "subcontracting"})
        # GeBIZ supplier — strengthens subcontracting evidence
        if ctx.get("gebiz_supplier"):
            backed.add("subcontracting")
        # SSL grade backs security measures and encryption
        if ssl_result and ssl_result.get("checked"):
            backed.add("security_measures")
            backed.add("encryption_standards")
        return backed

    def _build_docx(
        self,
        company_name: str,
        vendor_url: str,
        qa_answers: Dict[str, str],
        ctx: Dict,
        tx_hash: Optional[str] = None,
        product_type: str = "rfp_complete",
        intake: dict | None = None,
        coverage_summary: str | None = None,
    ) -> bytes:
        """Generate an editable DOCX evidence pack (Complete tier only)."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from io import BytesIO
            from app.core.config import settings

            doc = Document()

            # Title
            title = doc.add_heading("RFP Compliance Evidence Pack", level=0)
            title.runs[0].font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

            intake = intake or {}
            doc.add_paragraph(f"Prepared for: {company_name}")
            _tender_line = self._tender_attribution(intake)
            if _tender_line:
                doc.add_paragraph(f"Prepared in response to: {_tender_line}")
            doc.add_paragraph(f"Website: {vendor_url}")
            if coverage_summary:
                doc.add_paragraph(f"Coverage: {coverage_summary}")
            uen_val = ctx.get("uen") or "_______________________________"
            doc.add_paragraph(f"UEN (Singapore Business Registration No.): {uen_val}")
            if ctx.get("acra_name"):
                doc.add_paragraph(f"ACRA Registered Name: {ctx['acra_name']}")
            if intake.get("dpo_name") or intake.get("dpo_email"):
                dpo_parts = [x for x in [intake.get("dpo_name", ""), intake.get("dpo_email", "")] if x]
                doc.add_paragraph(f"DPO Contact: {', '.join(dpo_parts)}")
            if intake.get("iso_cert_number"):
                expiry = intake.get("iso_cert_expiry", "")
                cert_str = f"ISO 27001 Certificate: {intake['iso_cert_number']}"
                if expiry:
                    cert_str += f" (expires {expiry})"
                doc.add_paragraph(cert_str)
            if ctx.get("privacy_policy_url"):
                doc.add_paragraph(f"Privacy Policy: {ctx['privacy_policy_url']}")
            if ctx.get("gebiz_supplier"):
                count = ctx.get("gebiz_contracts_count") or 0
                doc.add_paragraph(
                    f"GeBIZ Registered Supplier — {count} prior government contract(s)" if count
                    else "GeBIZ Registered Supplier"
                )
            doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%d %b %Y %H:%M UTC')}")
            doc.add_paragraph(f"Report ID: {self.report_id}")
            if tx_hash:
                doc.add_paragraph(f"Blockchain TX: {tx_hash} ({settings.active_polygon_network_name})")
            doc.add_paragraph("")
            scope_para = doc.add_paragraph(
                "Scope of Assessment: This compliance pack is based on information provided by the "
                "company's authorised representative and automated website assessment conducted by "
                "Booppa on the date indicated."
            )
            scope_para.runs[0].italic = True
            doc.add_paragraph("")

            doc.add_heading("Compliance Q&A", level=1)
            for key, answer in qa_answers.items():
                doc.add_heading(self._q_label(key), level=2)
                p = doc.add_paragraph(answer if (answer and answer.strip()) else "Not provided.")
                if p.runs:
                    p.runs[0].font.size = Pt(11)
                doc.add_paragraph("")  # spacer

            # Attestation section
            doc.add_page_break()
            doc.add_heading("Vendor Attestation", level=1)
            doc.add_paragraph(
                f"I, the authorised representative of {company_name}, hereby attest that the "
                f"information provided in this RFP compliance evidence pack is accurate and "
                f"truthful to the best of my knowledge as of the date stated above."
            )
            doc.add_paragraph("\n\nSignature: _______________________________")
            doc.add_paragraph("Name: _______________________________")
            doc.add_paragraph("Designation: _______________________________")
            doc.add_paragraph("Date: _______________________________")
            doc.add_paragraph("Company Stamp (if applicable): _______________________________")

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            self.errors.append(f"DOCX error: {e}")
            return b""

    # ── Step 4: upload to S3 ──────────────────────────────────────────────────

    async def _write_certificate_log(self, pdf_bytes: bytes, download_url: str, db) -> None:
        """4.11: Write CertificateLog row for audit trail. Skipped for anonymous vendors."""
        import re as _re, hashlib as _hl
        if not db:
            return
        # CertificateLog.vendor_id is a UUID FK — skip if vendor_id is an email
        if not _re.match(r'^[0-9a-f-]{36}$', self.vendor_id or '', _re.IGNORECASE):
            return
        try:
            from app.core.models import CertificateLog
            import uuid as _uuid
            # Use the real S3 key set by _upload_pdf so downstream readers can
            # presign the URL. Falls back to the (legacy, wrong) express path
            # only if _upload_pdf wasn't reached for some reason.
            s3_key = getattr(self, "pdf_s3_key", None) or f"rfp-express/{self.report_id}.pdf"
            log = CertificateLog(
                vendor_id=_uuid.UUID(self.vendor_id),
                certificate_type="RFP",
                report_id=_uuid.UUID(self.report_id) if self.report_id else None,
                file_key=s3_key,
                file_hash=_hl.sha256(pdf_bytes).hexdigest() if pdf_bytes else None,
            )
            db.add(log)
            db.commit()
            logger.info(f"CertificateLog written for report {self.report_id}")
        except Exception as e:
            logger.warning(f"CertificateLog write failed (non-blocking): {e}")

    async def _upload_pdf(self, pdf_bytes: bytes, product_type: str = "rfp_express") -> str:
        try:
            from app.services.storage import S3Service
            s3 = S3Service()
            folder = "rfp-complete" if product_type == "rfp_complete" else "rfp-express"
            self.pdf_s3_key = f"reports/{folder}/{self.report_id}.pdf"
            url = await s3.upload_pdf(pdf_bytes, f"{folder}/{self.report_id}")
            return url
        except Exception as e:
            logger.error(f"S3 upload failed: {e}")
            self.errors.append(f"Upload error: {e}")
            raise

    async def _upload_docx(self, docx_bytes: bytes) -> Optional[str]:
        try:
            from app.services.storage import S3Service
            import boto3
            s3_svc = S3Service()
            key = f"rfp-complete/{self.report_id}.docx"
            s3_svc.s3_client.put_object(
                Bucket=s3_svc.bucket,
                Key=key,
                Body=docx_bytes,
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            url = s3_svc.s3_client.generate_presigned_url(
                "get_object",
                Params={"Bucket": s3_svc.bucket, "Key": key},
                ExpiresIn=7 * 24 * 3600,
            )
            logger.info(f"DOCX uploaded: {key}")
            return url
        except Exception as e:
            logger.error(f"DOCX upload failed: {e}")
            self.warnings.append(f"DOCX upload error: {e}")
            return None

    async def _upload_declaration(self, declaration_bytes: bytes) -> Optional[str]:
        try:
            from app.services.storage import S3Service
            s3_svc = S3Service()
            key = f"rfp-complete/{self.report_id}-declaration.pdf"
            s3_svc.s3_client.put_object(
                Bucket=s3_svc.bucket,
                Key=key,
                Body=declaration_bytes,
                ContentType="application/pdf",
            )
            url = s3_svc.s3_client.generate_presigned_url(
                "get_object",
                Params={"Bucket": s3_svc.bucket, "Key": key},
                ExpiresIn=7 * 24 * 3600,
            )
            logger.info(f"Supplier declaration uploaded: {key}")
            return url
        except Exception as e:
            logger.error(f"Declaration upload failed: {e}")
            self.warnings.append(f"Declaration upload error: {e}")
            return None

    async def _upload_appendix_d(self, appendix_bytes: bytes) -> Optional[str]:
        try:
            from app.services.storage import S3Service
            s3_svc = S3Service()
            key = f"rfp-complete/{self.report_id}-appendix-d.pdf"
            s3_svc.s3_client.put_object(
                Bucket=s3_svc.bucket,
                Key=key,
                Body=appendix_bytes,
                ContentType="application/pdf",
            )
            url = s3_svc.s3_client.generate_presigned_url(
                "get_object",
                Params={"Bucket": s3_svc.bucket, "Key": key},
                ExpiresIn=7 * 24 * 3600,
            )
            logger.info(f"Appendix D uploaded: {key}")
            return url
        except Exception as e:
            logger.error(f"Appendix D upload failed: {e}")
            self.warnings.append(f"Appendix D upload error: {e}")
            return None

    # ── Step 5: email ─────────────────────────────────────────────────────────

    async def _send_email(self, company_name: str, download_url: str, product_type: str = "rfp_express", docx_url: Optional[str] = None, declaration_url: Optional[str] = None, appendix_d_url: Optional[str] = None, pdf_bytes: Optional[bytes] = None):
        try:
            from app.services.rfp_express_emailer import RFPExpressEmailer
            emailer = RFPExpressEmailer()
            await emailer.send_express_ready_email(
                customer_email=self.vendor_email,
                vendor_name=company_name,
                download_url=download_url,
                product_type=product_type,
                declaration_url=declaration_url,
                appendix_d_url=appendix_d_url,
                docx_url=docx_url,
                pdf_bytes=pdf_bytes,
            )
        except Exception as e:
            logger.warning(f"Email delivery failed (non-blocking): {e}")
            self.warnings.append(f"Email not sent: {e}")
